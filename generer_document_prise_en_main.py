#!/usr/bin/env python3
"""Génère le document Word « Prise en main d'OpenVLA » avec schémas et contenu pédagogique."""

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCHEMA = ROOT / "assets" / "schema_openvla_infographie.png"
MACRO_DETAIL = ROOT / "assets" / "vue_macro_flux_donnees_detail.png"
FLUX_PNG = ROOT / "assets" / "flux_macro_openvla.png"
OUT = ROOT / "OpenVLA_prise_en_main.docx"

# Schéma ASCII (flux macro) — aligné sur la version fournie
FLUX_ASCII = r"""
                ┌──────────────────────────┐
                │      Image brute         │
                └─────────────┬────────────┘
                              │
                              ▼
                ┌──────────────────────────┐
                │     SigLIP Encoder       │
                │  (Extraction de patches) │
                └─────────────┬────────────┘
                              │
                              ▼
        ┌────────────────────────────────────────────────┐
        │            Framework Prismatic                 │
        │  - Orchestration du flux                       │
        │  - Injection des patches dans le flux texte    │
        └─────────────┬──────────────────────────────────┘
                      │
                      ▼
            ┌──────────────────────────┐
            │        Llama 2           │
            │ (Décodeur → prédit le    │
            │   prochain token d’action)│
            └─────────────┬────────────┘
                          │
                          ▼
            ┌──────────────────────────┐
            │       Action Head        │
            │ - Dé-quantification      │
            │ - Conversion tokens →    │
            │   commandes servo        │
            └─────────────┬────────────┘
                          │
                          ▼
                ┌──────────────────────────┐
                │   Bras robotique (Servo) │
                └──────────────────────────┘
""".strip("\n")


def render_flux_macro_png(path: Path) -> None:
    """Figure verticale : flux macro image → SigLIP → Prismatic → Llama → Action Head → robot."""
    mpl_dir = ROOT / ".matplotlib_cache"
    mpl_dir.mkdir(parents=True, exist_ok=True)
    os.environ["MPLCONFIGDIR"] = str(mpl_dir)

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    fig, ax = plt.subplots(figsize=(6.8, 8.6), dpi=160)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def add_box(cx: float, y0: float, w: float, h: float, lines: list[str], *, wide: bool = False) -> tuple[float, float]:
        x = cx - w / 2
        patch = FancyBboxPatch(
            (x, y0),
            w,
            h,
            boxstyle="round,pad=0.008,rounding_size=0.012",
            linewidth=1.35,
            edgecolor="#1a252f",
            facecolor="#eef2f5",
        )
        ax.add_patch(patch)
        yy = y0 + h - 0.022
        fs = 10.2 if not wide else 9.4
        for i, line in enumerate(lines):
            ax.text(cx, yy - i * 0.028, line, ha="center", va="top", fontsize=fs, color="#1a252f")
        return y0, y0 + h

    def arrow(y_from: float, y_to: float) -> None:
        arr = FancyArrowPatch(
            (0.5, y_from),
            (0.5, y_to),
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.2,
            color="#34495e",
            shrinkA=2,
            shrinkB=2,
        )
        ax.add_patch(arr)

    cx = 0.5
    # (y_bottom, height, width_factor 1.0 = 0.52, lines)
    stack: list[tuple[float, float, float, list[str]]] = [
        (0.86, 0.065, 0.52, ["Image brute"]),
        (0.755, 0.095, 0.52, ["SigLIP Encoder", "(Extraction de patches)"]),
        (0.58, 0.14, 0.78, ["Framework Prismatic", "— Orchestration du flux", "— Injection des patches dans le flux texte"]),
        (0.415, 0.11, 0.52, ["Llama 2", "(Décodeur → prochain token d’action)"]),
        (0.245, 0.12, 0.52, ["Action Head", "— Dé-quantification", "— Tokens → commandes servo"]),
        (0.045, 0.075, 0.52, ["Bras robotique (Servo)"]),
    ]

    tops: list[float] = []
    bottoms: list[float] = []
    y_cursor = 0.98
    for y0, h, w, lines in stack:
        y0 = y_cursor - h
        b, t = add_box(cx, y0, w, h, lines, wide=w > 0.6)
        bottoms.append(b)
        tops.append(t)
        y_cursor = y0 - 0.028

    for i in range(len(tops) - 1):
        arrow(bottoms[i] + 0.01, tops[i + 1] - 0.01)

    ax.text(
        0.5,
        0.995,
        "Vue macro du flux de données",
        ha="center",
        va="top",
        fontsize=13,
        fontweight="bold",
        color="#1a252f",
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_ascii_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def main() -> None:
    render_flux_macro_png(FLUX_PNG)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Prise en main d'OpenVLA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Schéma récapitulatif (vision, langage, action)", level=1)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(
        "Figure 1 — Architecture conceptuelle et flux de données (boucle VLA, entrées/sorties, "
        "empilement vision → adaptateur → LLM → tête d’action)."
    )
    r.italic = True

    if not SCHEMA.is_file():
        raise FileNotFoundError(f"Image introuvable : {SCHEMA}")
    doc.add_picture(str(SCHEMA), width=Inches(6.3))

    doc.add_heading("Lecture du schéma (panneaux A, B et C)", level=1)
    doc.add_paragraph(
        "Panneau A — Boucle VLA : on résume le problème comme une boucle Vision → Langage → "
        "Action. La vision capte l’état du monde, le langage fixe l’intention (consigne en "
        "langage naturel), l’action modifie le monde ; la caméra observe à nouveau l’effet, "
        "ce qui ferme la boucle pour des tâches séquentielles ou réactives."
    )
    doc.add_paragraph(
        "Panneau B — Entrées et sorties : une image (scène manipulable) et un prompt texte "
        "sont traités par un grand modèle type transformeur. La sortie n’est pas du texte "
        "conversationnel classique mais une séquence de jetons d’action, ensuite convertie "
        "en commandes de contrôle (souvent un contrôle continu discrétisé en petits pas pour "
        "le robot)."
    )
    doc.add_paragraph(
        "Panneau C — Découpage architectural : encodeur visuel (souvent type ViT / SigLIP), "
        "un projecteur MLP qui aligne les embeddings image sur l’espace d’entrée du modèle "
        "de langage, un « backbone » transformeur (ici Llama 2) qui fusionne texte et vision, "
        "puis une tête qui prédit les jetons d’action jusqu’au préhenseur."
    )

    doc.add_heading("Compréhension synthétique d’OpenVLA", level=1)
    doc.add_paragraph(
        "OpenVLA est un modèle vision-langage-action ouvert : il apprend à associer ce que "
        "voit une caméra, une consigne en langage naturel et une suite de petites actions "
        "robotiques. L’idée clé est de traiter le contrôle comme une génération de séquence "
        "dans un vocabulaire étendu (texte + jetons d’action), ce qui réutilise les mécanismes "
        "d’attention et d’auto-régression déjà efficaces pour le langage."
    )
    doc.add_paragraph(
        "Comparé à une chaîne classique perception → modèle géométrique → planificateur → "
        "contrôleur, une VLA compresse une partie du raisonnement et du réglage fin dans un "
        "seul réseau entraîné sur de grandes bases de démonstrations multi-robots, au prix "
        "d’exigences compute, de latence et de besoins de données pour un nouveau poste de travail."
    )

    doc.add_heading("Prismatic et Llama : deux rôles complémentaires", level=1)
    doc.add_paragraph(
        "Prismatic désigne ici la famille de VLM (vision-language model) pré-entraînée sur "
        "laquelle s’appuie OpenVLA — le dépôt Hugging Face « openvla-7b-prismatic » rappelle "
        "cette filiation. Prismatic couvre surtout la partie « comment aligner vision et langage » "
        "(encodeur(s) visuel(s), fusion, entraînement multimodal)."
    )
    doc.add_paragraph(
        "Llama 2 (7B) joue le rôle de moteur séquentiel : c’est le transformeur décodeur qui "
        "ingère le texte tokenisé et les représentations visuelles projetées, et qui prédit "
        "la suite — y compris les jetons d’action. En résumé : Prismatic ≈ socle VLM et recette "
        "d’alignement ; Llama ≈ cœur de raisonnement et de génération de la séquence de sortie."
    )

    doc.add_heading("Ordre de grandeur des poids", level=1)
    doc.add_paragraph(
        "Le checkpoint standard OpenVLA compte environ 7 milliards de paramètres (famille 7B). "
        "C’est le même ordre de grandeur que le backbone langage indiqué sur le schéma (Llama 2 7B), "
        "auquel s’ajoutent les paramètres de l’encodeur visuel, du projecteur et de la tête d’action. "
        "À titre de repère mentionné dans la littérature OpenVLA, ce format reste nettement plus "
        "léger que certains concurrents propriétaires beaucoup plus grands (par exemple des "
        "variantes autour de 55B pour d’autres approches VLA), tout en restant exigeant pour "
        "l’inférence temps réel sur GPU récent."
    )

    doc.add_heading("Architecture (vue détaillée)", level=1)
    doc.add_paragraph(
        "D’après les publications et matériels officiels sur OpenVLA, l’empilement typique est :"
    )
    items_arch = [
        "Encodeur visuel : combinaison SigLIP et DINOv2 (fusion de deux sources de représentations "
        "visuelles), plutôt qu’un simple SigLIP seul — le schéma pédagogique peut simplifier en "
        "« SigLIP / ViT ».",
        "Projecteur (MLP) : met les patchs / tokens visuels dans l’espace latent compatible avec "
        "les embeddings de mots du LLM.",
        "Backbone langage : Llama 2 7B avec mécanismes d’attention pour intégrer consigne et contexte visuel.",
        "Tête de prédiction des actions : étend le vocabulaire pour inclure des jetons d’action ; "
        "la sortie est décodée en incréments de contrôle (positions, orientations, gripper, etc., "
        "selon l’espace d’action du jeu de données).",
    ]
    for t in items_arch:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Vue macro du flux de données", level=1)
    doc.add_paragraph(
        "Du capteur à l’effecteur, en incluant la consigne utilisateur : l’instruction en langage "
        "naturel et l’image (capture RGB) alimentent l’encodeur SigLIP (patches / embeddings), le "
        "cadre Prismatic fusionne vision et tokens texte, Llama 2 prédit le prochain jeton d’action, "
        "la tête d’action dé-quantifie et convertit en commandes (ex. coordonnées articulaires), puis "
        "le bras exécute. Une boucle de retour (image, état robot, succès ou échec) peut réinjecter "
        "du contexte vers l’inférence pour la prochaine décision."
    )

    if not MACRO_DETAIL.is_file():
        raise FileNotFoundError(f"Image introuvable : {MACRO_DETAIL}")
    cap_detail = doc.add_paragraph()
    cap_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = cap_detail.add_run(
        "Figure 2 — Vue macro détaillée du flux de données (étapes 0 à 7 : instruction, capture, "
        "encodage SigLIP, Prismatic, décodeur Llama 2, Action Head, action physique, feedback)."
    )
    rd.italic = True
    doc.add_picture(str(MACRO_DETAIL), width=Inches(6.3))

    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = cap2.add_run(
        "Figure 3 — Schéma vertical simplifié (même enchaînement cœur modèle : image → SigLIP → "
        "Prismatic → Llama 2 → Action Head → robot)."
    )
    r2.italic = True

    doc.add_picture(str(FLUX_PNG), width=Inches(5.2))

    doc.add_paragraph("Équivalent en schéma texte (monospace) :")
    add_ascii_block(doc, FLUX_ASCII)

    doc.add_heading("Vue macro système (cellule robot)", level=1)
    doc.add_paragraph(
        "À l’échelle d’une cellule : capteurs (caméra 2D/3D, éventuellement état joint), prétraitement, "
        "inférence sur GPU, bus vers le contrôleur du robot (ex. UR + middleware), sécurité "
        "collaborative, puis retour capteur pour boucler la tâche. Le schéma ci-dessus se concentre "
        "sur le chemin logiciel à l’intérieur du modèle et du décodeur d’actions."
    )

    doc.add_heading("Exemple concret (comme sur le schéma)", level=1)
    doc.add_paragraph(
        "Scène : une table avec une sphère verte et un cube rouge. Consigne : « Ramasse le cube rouge ». "
        "L’image et le texte passent dans le modèle ; celui-ci émet une courte séquence de jetons "
        "d’action (symboles du type <act_k>). Un module aval traduit ces jetons en une trajectoire "
        "discrétisée dans l’espace opérationnel du robot (pas de translation/rotation, ouverture "
        "pince), exécutée par le bras jusqu’à la pose de préhension. Une nouvelle image confirme "
        "si la tâche a réussi ou s’il faut ré-inférer."
    )

    doc.add_heading("Sources et précision du document", level=1)
    doc.add_paragraph(
        "Ce document synthétise le schéma fourni et les descriptions publiques du projet OpenVLA "
        "(site du projet, article associé, fiches modèle). Les détails d’implémentation (taille "
        "exacte des images, normalisation, espace d’action du robot cible) doivent être vérifiés "
        "dans la version du code et du checkpoint utilisés pour votre démonstrateur."
    )

    doc.save(OUT)
    print(f"Figure flux : {FLUX_PNG}")
    print(f"Document créé : {OUT}")


if __name__ == "__main__":
    main()
