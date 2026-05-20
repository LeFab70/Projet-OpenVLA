#!/usr/bin/env python3
"""Génère le document Word de description de stage OpenVLA."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("openVLA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Intro stage", level=1)
    p_intro = doc.add_paragraph(
        "Développement d’un démonstrateur OpenVLA pour robot collaboratif UR "
        "et vision 2D ou 3D ZIVID"
    )
    p_intro.paragraph_format.space_after = Pt(12)

    doc.add_heading("Objectifs", level=1)
    doc.add_paragraph(
        "Le projet visera en particulier à répondre à la question suivante : "
        "quels sont les ingrédients de la « sauce OpenVLA » et en quoi cette "
        "approche diffère-t-elle des méthodes robotiques plus traditionnelles ? "
        "Le stage permettra donc de combiner veille technologique, expérimentation "
        "pratique, intégration robotique et vulgarisation technique."
    )

    doc.add_heading("Principales responsabilités", level=1)
    responsabilites = [
        "Étudier l’architecture générale d’OpenVLA et en produire une description "
        "technique claire et structurée.",
        "Identifier les différentes composantes de la pipeline : perception, "
        "représentation des données, traitement des entrées multimodales, inférence, "
        "génération d’actions et exécution robotique.",
        "Comparer l’approche OpenVLA aux approches robotiques traditionnelles, "
        "notamment en ce qui concerne la programmation de tâches, la perception, "
        "la planification et le contrôle.",
        "Mettre en place un démonstrateur expérimental utilisant un robot "
        "collaboratif UR et une caméra Zivid.",
        "Réaliser l’intégration logicielle et matérielle nécessaire au "
        "fonctionnement du démonstrateur.",
        "Effectuer des essais, analyser les résultats obtenus et documenter les "
        "limites, défis et opportunités de l’approche.",
        "Produire une documentation technique et une synthèse vulgarisée destinée "
        "à l’équipe de CCNB-INNOV.",
        "Présenter les résultats du stage sous forme de démonstration et de "
        "présentation technique.",
    ]
    for item in responsabilites:
        doc.add_paragraph(item, style="List Bullet")

    out = Path(__file__).resolve().parent / "openVLA_stage_CCNB.docx"
    doc.save(out)
    print(f"Document créé : {out}")


if __name__ == "__main__":
    main()
