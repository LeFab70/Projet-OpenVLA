#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère agenda_quotidien_Fabrice.docx (structure alignée rapports_hebdomadaires)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DIR = Path(__file__).resolve().parent
OUT_DOCX = DIR / "agenda_quotidien_Fabrice.docx"
OUT_MD = DIR / "agenda_quotidien_Fabrice.md"

SUPERVISEUR = "Guillaume Batungwanayo"
STAGIAIRE = "Fabrice Kouonang"

# Chaque jour : titre, date, objectif, realisations[], problemes[], livrables[], prochaine_etape
JOURS = [
    {
        "num": "01",
        "titre": "Début de stage CCNB-INNOV",
        "date": "14 mai 2026 (jeudi)",
        "objectif": (
            "Démarrer le stage au CCNB-INNOV, cadrer le sujet du démonstrateur OpenVLA "
            "(robot UR + caméra Zivid 2+ MR130) et lancer le rapport final."
        ),
        "realisations": [
            "Présentation du stage, des objectifs et des responsabilités (veille, intégration, démonstration).",
            "Formulation de la question directrice : quels sont les ingrédients de la « sauce OpenVLA » ?",
            "Structuration du rapport final : parties I (architecture), II (pipeline Zivid / OpenVLA / UR), III (comparaison).",
            "Première description technique d’OpenVLA et du positionnement par rapport à la robotique traditionnelle.",
        ],
        "problemes": [],
        "livrables": [
            "OpenVLA_day01_stage_CCNB.docx — parties I à III, figures 1–4.",
            "README du dépôt — journal de bord (Jour 01).",
        ],
        "prochaine": "Étudier l’architecture OpenVLA en détail et installer openvla-7b sur le poste Windows.",
    },
    {
        "num": "02",
        "titre": "Prise en main OpenVLA",
        "date": "15 mai 2026 (vendredi)",
        "objectif": (
            "Comprendre la boucle Vision–Langage–Action (VLA) et mettre en place l’environnement "
            "d’inférence openvla-7b sous Windows 11 Pro."
        ),
        "realisations": [
            "Veille : encodeur visuel (SigLIP / DINOv2), projecteur MLP, LLM Llama 2 (> 95 % des poids).",
            "Installation Conda env_openvla (Python 3.11), transformers==4.40.1, PyTorch CUDA, Hugging Face.",
            "Téléchargement et test de chargement du modèle openvla-7b (~15 Go VRAM en FP16).",
            "Comparaison conceptuelle OpenVLA vs robotique classique (langage, stratégie, adaptation).",
            "Documentation des commandes et chemins dans scripts/utils.txt.",
        ],
        "problemes": [
            "GPU RTX 5090 (Blackwell sm_120) non supporté par PyTorch stable → PyTorch nightly CUDA 12.8 (cu128).",
            "Version transformers stricte pour OpenVLA → figer transformers==4.40.1.",
            "Modèle volumineux → chargement float16/bfloat16 sur GPU ≥ 16 Go.",
        ],
        "livrables": [
            "OpenVLA_day02_prise_en_main.docx",
            "scripts/integration/test/test_openvla.py (vérification GPU / VRAM)",
            "scripts/utils.txt",
        ],
        "prochaine": "Configurer Conda (Zivid, UR), prendre en main le robot UR et l’API Zivid.",
    },
    {
        "num": "03",
        "titre": "Robot UR — tracé lettre A",
        "date": "19 mai 2026 (mardi)",
        "objectif": "Premiers mouvements programmés sur le bras collaboratif UR via URScript (approche sécurisée, tracé géométrique).",
        "realisations": [
            "Définition de 5 poses clés pour tracer la lettre « A » (movej approche/retrait, movel traits linéaires).",
            "Paramètres de sécurité : vitesse ~0,1 m/s, lift 0,02 m entre segments.",
            "Scripts URscriptLetterA.script (boucle 1–4 lettres), traceAOnce.script, returnToCenter.script.",
            "Essai complémentaire sur PolyScope (saisie point par point, sans script complet).",
        ],
        "problemes": [
            "Première prise en main du TCP et des repères → validation pose par pose sur le teach pendant.",
        ],
        "livrables": [
            "OpenVLA_day03_trace_A.docx",
            "scripts/ur/*.script (versionnés Git)",
        ],
        "prochaine": "Connecter la caméra Zivid MR130 et isoler les environnements Python.",
    },
    {
        "num": "04",
        "titre": "API Zivid + environnements Conda",
        "date": "20 mai 2026 (mercredi)",
        "objectif": "Maîtriser la capture 2D/3D Zivid et structurer les environnements Conda par brique logicielle.",
        "realisations": [
            "Connexion Zivid 2+ MR130, acquisition capture_2d_3d.",
            "Script scripts/zivid/capture.py : export ColorImage.png, Frame.zdf, PointCloud.ply.",
            "Création de 3 environnements Conda (Python 3.11) : env_zivid, env_ur, env_integration.",
            "Guide Conda/Anaconda : arborescence, activation, dépendances par module.",
        ],
        "problemes": [
            "Plage de distance MR130 : profondeur bruitée hors zone de travail → respect distance fabricant, bornes Z à la capture.",
        ],
        "livrables": [
            "OpenVLA_day04_zivid_api.docx",
            "OpenVLA_day04_conda_anaconda.docx",
            "scripts/zivid/capture.py",
        ],
        "prochaine": "Chaîner capture Zivid et inférence OpenVLA sur GPU.",
    },
    {
        "num": "05",
        "titre": "OpenVLA + intégration Zivid",
        "date": "21 mai 2026 (jeudi)",
        "objectif": "Valider la pipeline perception (Zivid) → inférence OpenVLA → prédiction d’actions 7D.",
        "realisations": [
            "Capture Zivid réussie ; conversion RGBA → RGB 224×224 pour l’encodeur visuel.",
            "Chargement openvla-7b sur GPU ; predict_action (XYZ, rotation, pince).",
            "Scripts test_openvla.py et test_zivid_openvla.py (consignes texte libres).",
            "Schéma d’intégration Zivid (SDK) + OpenVLA (transformers) ; connexion UR-RTDE planifiée.",
        ],
        "problemes": [
            "OpenVLA impossible sur Mac (pas de CUDA) → déploiement PC Windows 11 Pro dédié.",
            "Conflits de dépendances Zivid / UR / OpenVLA → environnements Conda séparés.",
        ],
        "livrables": [
            "OpenVLA_day05_openvla_integration.docx",
            "scripts/integration/test/test_openvla.py",
            "scripts/integration/test/test_zivid_openvla.py",
        ],
        "prochaine": "Fermer la boucle avec le UR16e (RTDE) et consolider le rapport final.",
    },
    {
        "num": "06",
        "titre": "Rapport final + démonstrateur UR/Zivid/OpenVLA",
        "date": "22–23 mai 2026 (vendredi–samedi labo)",
        "objectif": "Consolider le rapport final (I–III) puis réaliser le premier démonstrateur en boucle fermée sur le robot.",
        "realisations": [
            "22 mai : édition et consolidation des parties I à III de OpenVLA_day01_stage_CCNB.docx.",
            "23 mai : boucle Zivid → OpenVLA → UR16e via demoTest.py (RTDE moveL, lecture TCP, pince digitale).",
            "Modes SAFE_MODE (simulation des poses) et SCALE max 5 cm/step ; vitesses et accélérations limitées.",
            "Documentation de l’arborescence scripts/ et des 4 environnements Conda (zivid, ur, integration, openvla).",
        ],
        "problemes": [
            "Risque de mouvements non maîtrisés → SAFE_MODE, SCALE=0,05, vitesses réduites avant essais réels.",
        ],
        "livrables": [
            "OpenVLA_day06_demo_ur_zivid.docx",
            "scripts/integration/testUR_ZIVID/demoTest.py",
            "OpenVLA_day01_stage_CCNB.docx (mise à jour 22 mai)",
        ],
        "prochaine": "Ajouter détection 2D (YOLO / Grounding DINO) et projection 3D dans la pipeline.",
    },
    {
        "num": "07",
        "titre": "Données robot, YOLO, Grounding DINO",
        "date": "25 mai 2026 (lundi)",
        "objectif": "Clarifier le flux Zivid → détection 2D → OpenVLA → UR et comparer YOLOv8n vs Grounding DINO.",
        "realisations": [
            "Pipeline commun : RGB + nuage → (u,v) → (X,Y,Z) → prompt OpenVLA → predict_action → RTDE.",
            "YOLOv8n : classes COCO, coordonnées XYZ injectées dans le prompt.",
            "Grounding DINO : open-vocabulary (ex. cell phone.), projection 3D pour le contrôleur.",
            "Scripts zivid_yolo_openvla.py, test_zivid_groundingDino.py, returnAllPositions.py.",
            "Mise à jour rapport final : sections II.1.1, II.1.2, II.1.3.",
        ],
        "problemes": [
            "Résolutions 2D ≠ 3D Zivid → attention au scale_u / scale_v pour la projection (préparation jour 10).",
        ],
        "livrables": [
            "OpenVLA_day07_robot_data.docx",
            "scripts/openVLA_ZIVID/test/*.py",
            "yolov8n.pt",
        ],
        "prochaine": "Boucle continue adaptative (réinférence à chaque nouvelle image).",
    },
    {
        "num": "08",
        "titre": "Boucle continue (OpenVLA adaptatif)",
        "date": "26 mai 2026 (mardi)",
        "objectif": "Tester une boucle perception → action → réinférence pour correction itérative avec Grounding DINO.",
        "realisations": [
            "Scripts demo_adaptatif_openvla.py (DINO + UR + boucle) et demo_adaptatif_openvla_print_value.py (simulation, logs).",
            "Injection optionnelle des coordonnées (X,Y,Z) DINO + Zivid dans le prompt (comme variante YOLO).",
            "Mise à jour rapport final II.5 (boucle continue + injection XYZ).",
        ],
        "problemes": [
            "Premiers signes d’oscillation possible → prévoir SCALE réduit et filtrage profondeur (jour 09–10).",
        ],
        "livrables": [
            "OpenVLA_day08_boucle_continue.docx",
            "scripts/integration/test/demo_adaptatif_openvla*.py",
        ],
        "prochaine": "Collecter métriques, interpréter convergence / divergence, documenter résultats.",
    },
    {
        "num": "09",
        "titre": "Tests, interprétation et rapports",
        "date": "27 mai 2026 (mercredi)",
        "objectif": "Valider et analyser la boucle adaptative ; interpréter l’impact des coordonnées DINO et de la projection 3D.",
        "realisations": [
            "Procédure : demo_adaptatif_openvla_print_value.py (itération, XYZ DINO, delta OpenVLA, TCP, distance, pince).",
            "Essais robot avec SAFE_MODE via demo_adaptatif_openvla.py.",
            "Métriques : distance euclidienne d_t, itérations jusqu’à convergence, stabilité Z, multi-détections DINO.",
            "Interprétation : convergence monotone, offset constant (calibration TCP), divergence (SCALE / détection).",
            "Appendice II.5 du rapport final avec résultats.",
        ],
        "problemes": [
            "Bruit Z sur nuage Zivid → moyenne locale / filtre profondeur envisagés.",
            "Oscillation autour d’un offset → hypothèse calibration TCP / transform (traité jour 11).",
        ],
        "livrables": [
            "OpenVLA_day09_boucle_results.docx",
            "OpenVLA_day01_stage_CCNB.docx (II.5)",
        ],
        "prochaine": "Refactoriser en modules pipeline/ et sécuriser workspace + calibration.",
    },
    {
        "num": "10",
        "titre": "Refactoring pipeline (modules indépendants)",
        "date": "28 mai 2026 (jeudi)",
        "objectif": "Découper le démonstrateur en modules réutilisables ; corriger bugs majeurs (calibration, workspace, DINO).",
        "realisations": [
            "Dossier pipeline/ : config, calibration, zivid_capture, dino_detector, ur_controller, gripper, vla_controller.",
            "main_real.py (robot réel), main_sim.py (dry-run), calibrer_robot.py (génération T_tcp_cam.npy).",
            "Règles : calibration main-œil obligatoire, bornes workspace, réintégration DINO toutes les N étapes.",
            "Prompt dynamique basé sur distance TCP→objet (deltas) plutôt que coordonnées absolues.",
        ],
        "problemes": [
            "IndexError dino_detector (u,v) hors nuage → scale_u/scale_v + clipping des indices.",
            "Dérive boucle continue → réinférence périodique DINO + réduction SCALE.",
        ],
        "livrables": [
            "OpenVLA_day10_refactoring_pipeline.docx",
            "pipeline/*.py",
        ],
        "prochaine": "Calibration eye-in-hand et rapport hebdomadaire semaine 3.",
    },
    {
        "num": "11",
        "titre": "Rapport hebdomadaire + calibration main-œil",
        "date": "29 mai 2026 (vendredi)",
        "objectif": "Produire le rapport hebdomadaire (semaine 26–29 mai) et générer T_tcp_cam.npy (hand-eye Zivid).",
        "realisations": [
            "Rapport Semaine_03_25-29_mai_2026.docx dans rapports_hebdomadaires/.",
            "calibrer_robot.py : poses variées, mire Zivid, detect_feature_points, calibrate_eye_in_hand, np.save.",
            "calibration.py : load_calibration(), cam_to_robot(), compute_distance_tcp_to_object().",
            "Règles de pose : inclinaison ±15–25°, hauteur Z 40–60 cm, rotation RZ.",
            "Rapport final : section II.6 Jour 11.",
        ],
        "problemes": [
            "Sans T_tcp_cam, repère caméra ≠ repère robot → script autonome de calibration avant main_real.",
        ],
        "livrables": [
            "OpenVLA_day11_rapport_hebdomadaire.docx",
            "rapports_hebdomadaires/Semaine_03_25-29_mai_2026.docx",
            "pipeline/calibrer_robot.py, pipeline/calibration.py",
        ],
        "prochaine": "Tests réels du démonstrateur avec calibration et validation workspace.",
    },
    {
        "num": "12",
        "titre": "Test réel du démonstrateur Zivid / OpenVLA",
        "date": "1 juin 2026 (lundi)",
        "objectif": (
            "Exécuter et valider la chaîne complète sur la cellule UR16e : capture Zivid, inférence OpenVLA, "
            "mouvements RTDE avec calibration T_tcp_cam et bornes de sécurité."
        ),
        "realisations": [
            "(Prévu) Lancer test_zivid_openvla.py — validation capture + inférence sans robot.",
            "(Prévu) Lancer demoTest.py puis pipeline/main_real.py avec T_tcp_cam.npy chargé.",
            "(Prévu) Vérifier SAFE_MODE, SCALE, workspace ; noter distances, convergence pince Robotiq.",
            "(Prévu) Consigner résultats pour rapport et agenda / semaine 4.",
        ],
        "problemes": [
            "(À compléter après essais) — calibration, profondeur MR130, stabilité DINO en conditions réelles.",
        ],
        "livrables": [
            "Scripts : test_zivid_openvla.py, demoTest.py, pipeline/main_real.py",
            "Fichier calibration : T_tcp_cam.npy (si généré)",
            "Mise à jour agenda + rapport journalier (à rédiger)",
        ],
        "prochaine": "Prise d’objet en boucle fermée (DINO + OpenVLA + pince) et poursuite documentation.",
    },
]


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _add_bullets(doc, items):
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def _add_section(doc, title, content_paragraph=None, bullets=None):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    if content_paragraph:
        doc.add_paragraph(content_paragraph)
    if bullets:
        _add_bullets(doc, bullets)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Projet OpenVLA — Agenda quotidien", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Par : {STAGIAIRE}")
    doc.add_paragraph("—" * 40)
    doc.add_paragraph(f"Superviseur direct : {SUPERVISEUR}")
    doc.add_paragraph("Période : 14 mai au 1 juin 2026 | Jours 01–12")
    doc.add_paragraph(
        "Document de suivi journalier (contenu aligné sur les rapports hebdomadaires "
        "et le README du dépôt Git)."
    )
    doc.add_page_break()

    for i, jour in enumerate(JOURS):
        _add_heading(doc, f"Jour {jour['num']} — {jour['titre']}", level=1)
        doc.add_paragraph(f"Date : {jour['date']}")

        _add_section(doc, "1. Objectif du jour", jour["objectif"])
        _add_section(doc, "2. Réalisations", bullets=jour["realisations"])

        if jour["problemes"]:
            _add_section(
                doc,
                "3. Problèmes rencontrés et solutions apportées",
                bullets=jour["problemes"],
            )
        else:
            _add_section(doc, "3. Problèmes rencontrés et solutions apportées")
            doc.add_paragraph("Aucun blocage majeur documenté pour cette journée.")

        _add_section(doc, "4. Livrables", bullets=jour["livrables"])
        _add_section(doc, "5. Prochaine étape", jour["prochaine"])
        if i < len(JOURS) - 1:
            doc.add_page_break()

    doc.add_paragraph()
    _add_heading(doc, "Repères calendaires", level=1)
    _add_bullets(
        doc,
        [
            "16–18 mai 2026 : pause / hors laboratoire (pas d’entrée journalière).",
            "24 mai 2026 : dimanche.",
            "30–31 mai 2026 : week-end.",
            "Semaines : Semaine_01 (14–15 mai), Semaine_02 (19–22 mai), Semaine_03 (25–29 mai) — voir rapports_hebdomadaires/.",
        ],
    )

    doc.save(OUT_DOCX)
    print(f"Écrit : {OUT_DOCX}")


def build_md():
    lines = [
        "# Agenda quotidien — Stage OpenVLA",
        "",
        f"**Programmeur :** {STAGIAIRE}  ",
        "**Début du stage :** 14 mai 2026  ",
        "**Document Word :** `agenda_quotidien_Fabrice.docx` (généré par ce script)",
        "",
        f"**Superviseur :** {SUPERVISEUR}",
        "",
        "---",
        "",
    ]
    for jour in JOURS:
        lines += [
            f"## Jour {jour['num']} — {jour['titre']}",
            "",
            f"**Date :** {jour['date']}",
            "",
            "### 1. Objectif du jour",
            "",
            jour["objectif"],
            "",
            "### 2. Réalisations",
            "",
        ]
        lines += [f"- {r}" for r in jour["realisations"]]
        lines += ["", "### 3. Problèmes et solutions", ""]
        if jour["problemes"]:
            lines += [f"- {p}" for p in jour["problemes"]]
        else:
            lines.append("- Aucun blocage majeur documenté.")
        lines += ["", "### 4. Livrables", ""]
        lines += [f"- {l}" for l in jour["livrables"]]
        lines += [
            "",
            "### 5. Prochaine étape",
            "",
            jour["prochaine"],
            "",
            "---",
            "",
        ]
    lines += [
        "## Repères calendaires",
        "",
        "- 16–18 mai : pause / hors labo",
        "- 24 mai : dimanche | 30–31 mai : week-end",
        "- Rapports hebdo : `rapports_hebdomadaires/Semaine_01` à `Semaine_03`",
        "",
    ]
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"Écrit : {OUT_MD}")


if __name__ == "__main__":
    build()
    build_md()
