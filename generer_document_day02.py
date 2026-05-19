#!/usr/bin/env python3
"""Génère le document Word — Jour 02 : tracé du A (URScript)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCRIPT_UR = ROOT / "scripts" / "ur" / "URscriptLetterA.script"
OUT = ROOT / "OpenVLA_day02_trace_A.docx"


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def main() -> None:
    if not SCRIPT_UR.is_file():
        raise FileNotFoundError(f"Script UR introuvable : {SCRIPT_UR}")
    trace_code = SCRIPT_UR.read_text(encoding="utf-8")
    script_rel = "scripts/ur/URscriptLetterA.script"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Projet OpenVLA — Jour 02", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Objectif de la journée", level=1)
    doc.add_paragraph(
        "Mise en place et validation d’un premier programme de contrôle sur le robot "
        "collaboratif UR : tracé géométrique de la lettre « A » dans l’espace de travail, "
        "à l’aide d’un script URScript (fonction trace_A avec répétition paramétrable)."
    )
    doc.add_paragraph(
        "Cette étape pose les bases de l’intégration robotique du stage : définition des poses, "
        "paramètres de mouvement (vitesse, accélération, lissage), séquence d’approche et "
        "retrait sécurisé de l’outil."
    )

    doc.add_heading("Travail réalisé", level=1)
    items = [
        "Configuration TCP et charge utile (set_tcp, set_payload).",
        "Définition des paramètres de mouvement : vitesse 0,1 m/s, accélération 0,5 m/s², "
        "rayon de lissage (blend) 0,005 m.",
        "Calcul des cinq points clés du A (base, sommet, barre transversale) dans le repère outil.",
        "Séquence movej / movel : approche, traits gauche et droit, levée d’outil, barre horizontale, "
        "retour en position d’approche.",
        "Boucle while (i < n) avec sync() : plusieurs lettres A, décalage sp sur l’axe Y.",
        "Contrôle de sécurité : n entre 1 et 4 (popup + halt).",
        "Une seule fonction trace_A(n) — ex. trace_A(1) en bas du fichier pour lancer.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Fichier script externe", level=1)
    doc.add_paragraph(
        f"Le programme URScript est maintenu dans un fichier séparé du dépôt, pour pouvoir "
        f"être chargé directement sur le robot ou dans l’émulateur UR sans repasser par le "
        f"document Word :"
    )
    p_path = doc.add_paragraph()
    p_path.add_run(script_rel).bold = True
    doc.add_paragraph(
        "Sur PolyScope : Programmes → Nouveau programme → coller le contenu du fichier, "
        "ou importer via l’éditeur de script. En simulation (URSim), ouvrir le même fichier "
        "dans l’éditeur intégré."
    )

    doc.add_heading("Programme URScript — URscriptLetterA.script (copie de référence)", level=1)
    doc.add_paragraph(
        "Reproduction du contenu du fichier externe ci-dessous (à jour si le fichier "
        f"{script_rel} est modifié, régénérer ce document avec generer_document_day02.py)."
    )
    add_code_block(doc, trace_code.rstrip() + "\n")

    doc.add_heading("Points d’attention", level=1)
    notes = [
        "Vérifier les limites de l’espace de travail et les zones interdites avant exécution.",
        "Ajuster bx, by, bz, w (largeur), h (hauteur) pour le poste UR + surface de tracé.",
        "rx = 3,14159 : outil vers le bas ; adapter selon le TCP.",
        "Espacement entre lettres : sp = 0,15 m (y = by + sp × i).",
        "Barre du A à 45 % de la hauteur (zb = bz + h × 0,45).",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Programmeur : ").bold = True
    meta.add_run("Fabrice Kouonang\n")
    meta.add_run("Projet : ").bold = True
    meta.add_run("OpenVLA — Stage CCNB-INNOV\n")
    meta.add_run("Jour : ").bold = True
    meta.add_run("02")

    doc.save(OUT)
    print(f"Script UR : {SCRIPT_UR}")
    print(f"Document créé : {OUT}")


if __name__ == "__main__":
    main()
