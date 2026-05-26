#!/usr/bin/env python3
"""Construit le rapport Jour 08 et met à jour le rapport final (boucle continue).

Jour 08 (26 mai 2026) : tests boucle continue perception→action→réinférence (OpenVLA adaptatif),
et injection optionnelle des coordonnées (X,Y,Z) issues de Grounding DINO dans l'instruction.
"""

import copy
import os
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(BASE, "OpenVLA_day05_openvla_integration.docx")
DAY08 = os.path.join(BASE, "OpenVLA_day08_boucle_continue.docx")
DAY01 = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")

MARKER_DAY08_FINAL = "Jour 08 — boucle continue perception → action → réinférence"


def logo_paragraph_element():
    doc5 = Document(TEMPLATE)
    return copy.deepcopy(doc5.paragraphs[0]._element)


def clear_body_keep_logo(doc: Document) -> None:
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def insert_after(elem, doc, text, style=None):
    new_p = OxmlElement("w:p")
    elem.addnext(new_p)
    para = Paragraph(new_p, doc)
    if style:
        para.style = style
    para.add_run(text)
    return new_p


def remove_paragraphs_containing(doc, needles):
    for p in list(doc.paragraphs):
        if any(n in p.text for n in needles):
            p._element.getparent().remove(p._element)


def build_day08() -> None:
    shutil.copy2(TEMPLATE, DAY08)
    doc = Document(DAY08)
    clear_body_keep_logo(doc)

    add_para(doc, "Projet OpenVLA — Jour 08 — Boucle continue perception → action → réinférence")
    add_para(doc, "Par : Fabrice Kouonang")
    add_para(doc, "————————————————————————————")
    add_para(doc, "Superviseur direct : Guillaume Batungwanayo")
    add_para(doc, "Date : 26 mai 2026 | Jour : 08")

    add_heading(doc, "Objectif de la journée")
    add_para(
        doc,
        "Tester une boucle continue (closed-loop) afin de rendre OpenVLA plus adaptatif : "
        "perception → action → nouvelle perception → réinférence. L'idée est d'exécuter plusieurs "
        "petites actions successives, plutôt qu'une action unique, en se réajustant à chaque "
        "nouvelle observation de la scène.",
    )

    add_heading(doc, "Contexte")
    add_para(
        doc,
        "Les jours précédents ont validé la chaîne Zivid → OpenVLA → UR16e. Le jour 07 a ajouté "
        "une détection amont (YOLO / Grounding DINO) et la projection 3D (X,Y,Z). Le jour 08 "
        "met l'accent sur la boucle continue pour améliorer la robustesse face aux imprévus "
        "(erreurs de trajectoire, obstacles, variations de scène).",
    )

    add_heading(doc, "1. Principe — boucle continue perception → action → réinférence")
    add_bullet(doc, "À chaque itération : capturer l'image Zivid (RGB 224×224 pour OpenVLA).")
    add_bullet(doc, "Inférer une action OpenVLA (delta 7D : translation, rotation, gripper).")
    add_bullet(doc, "Exécuter une petite étape (SCALE limité) puis recapturer et réinférer.")
    add_bullet(doc, "Condition d'arrêt : objet saisi (pince fermée / seuil) ou nombre d'étapes max.")

    add_heading(doc, "2. Implémentation testée (scripts jour 08)")
    add_para(
        doc,
        "Deux scripts ont été ajoutés pour tester le comportement adaptatif, en deux phases :",
    )
    add_bullet(doc, "Phase 1 : Grounding DINO localise l'objet demandé (texte libre) et récupère (X,Y,Z) via le nuage Zivid.")
    add_bullet(doc, "Phase 2 : OpenVLA boucle perception→action→réinférence jusqu'à prise de l'objet (ou arrêt sécurité).")

    add_heading(doc, "2.1 Phase 1 — Grounding DINO + coordonnées 3D", level=2)
    add_para(
        doc,
        "Grounding DINO reçoit une requête textuelle (ex. « cell phone. ») et retourne une boîte 2D. "
        "Le centre pixel (u,v) est projeté dans le nuage de points Zivid afin d'obtenir (X,Y,Z) en mètres.",
    )

    add_heading(doc, "2.2 Phase 2 — OpenVLA adaptatif (réinférence continue)", level=2)
    add_para(
        doc,
        "Au lieu d'une action unique, OpenVLA prédit une petite correction itérative. Cette approche "
        "rend le système plus robuste : si la scène change ou si l'action n'est pas parfaite, la "
        "réinférence corrige au pas suivant.",
    )

    add_heading(doc, "3. Prompt OpenVLA — injection optionnelle des XYZ (efficacité)")
    add_para(
        doc,
        "Observation importante : même si le test Grounding DINO jour 07 utilisait un prompt sans XYZ "
        "(label seul), il est possible d'injecter les coordonnées (X,Y,Z) issues de Grounding DINO "
        "dans l'instruction envoyée à OpenVLA. Cela peut guider davantage l'action et améliorer "
        "l'efficacité, surtout quand l'objet est petit ou partiellement occulté.",
    )
    add_para(doc, "Prompt minimal (label seul) :", bold=True)
    add_para(doc, "pick up the {label}", bold=False)
    add_para(doc, "Prompt enrichi (label + XYZ projetés) :", bold=True)
    add_para(doc, "the {label} is located at position X=…m Y=…m Z=…m. What action should the robot take to pick it up?", bold=False)

    add_heading(doc, "4. Sécurité et paramètres")
    add_bullet(doc, "SCALE limité pour des mouvements incrémentaux (sécurité).")
    add_bullet(doc, "MAX_STEPS pour éviter les boucles infinies.")
    add_bullet(doc, "SAFE_MODE pour valider la logique sans mouvement réel du robot.")
    add_bullet(doc, "Seuil pince (gripper) pour décider l'arrêt (objet saisi).")

    add_heading(doc, "5. Scripts de référence (jour 08)")
    add_bullet(doc, "scripts/integration/test/demo_adaptatif_openvla.py — DINO + UR16e + boucle OpenVLA (SAFE_MODE possible).")
    add_bullet(doc, "scripts/integration/test/demo_adaptatif_openvla_print_value.py — DINO + boucle OpenVLA (simulation, affichage valeurs).")

    add_heading(doc, "Livrables jour 08")
    add_bullet(doc, "OpenVLA_day08_boucle_continue.docx (ce document).")
    add_bullet(doc, "Rapport final : ajout d'une note sur la boucle continue et l'injection XYZ (Grounding DINO).")
    add_bullet(doc, "README — entrée Jour 08 mise à jour.")

    doc.save(DAY08)
    print(f"Built {DAY08}")


def update_day01_for_day08() -> None:
    doc = Document(DAY01)

    remove_paragraphs_containing(
        doc,
        [
            MARKER_DAY08_FINAL,
            "demo_adaptatif_openvla.py",
            "demo_adaptatif_openvla_print_value.py",
            "réinférence continue",
            "reinférence continue",
        ],
    )

    # Insérer une note dans II.5 (après la liste d'entrées modèle si présente)
    inserted = False
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("II.5"):
            anchor = p._element
            # chercher la puce "Entrées modèle :" juste après, sinon insérer juste après le titre
            for p2 in doc.paragraphs:
                if p2.text.strip().startswith("Entrées modèle : image RGB") or p2.text.strip().startswith("Entrées modèle: image RGB"):
                    anchor = p2._element
                    break

            anchor = insert_after(anchor, doc, MARKER_DAY08_FINAL, "Heading 3")
            anchor = insert_after(
                anchor,
                doc,
                "Jour 08 : test d'une boucle continue perception → action → réinférence pour rendre "
                "le comportement plus adaptatif (correction itérative à chaque nouvelle observation).",
                "Normal",
            )
            anchor = insert_after(
                anchor,
                doc,
                "Option d'efficacité : injecter dans l'instruction OpenVLA les coordonnées (X,Y,Z) "
                "projetées issues de Grounding DINO + nuage Zivid (en plus du label), comme la variante YOLO.",
                "Normal",
            )
            insert_after(anchor, doc, "Scripts : demo_adaptatif_openvla.py, demo_adaptatif_openvla_print_value.py.", "List Bullet")
            inserted = True
            break

    if not inserted:
        # fallback : juste avant III
        for p in doc.paragraphs:
            if p.style.name == "Heading 1" and p.text.strip().startswith("III."):
                anchor = p._element.getprevious() or p._element
                insert_after(anchor, doc, MARKER_DAY08_FINAL, "Heading 3")
                insert_after(
                    anchor,
                    doc,
                    "Jour 08 : test boucle continue perception → action → réinférence (OpenVLA adaptatif) "
                    "avec option d'injection XYZ depuis Grounding DINO.",
                    "Normal",
                )
                break

    # Ajuster la puce II.5 entrées modèle si présente
    for p in doc.paragraphs:
        if p.text.strip().startswith("Entrées modèle : image RGB 224×224"):
            if "Grounding DINO" in p.text and "label seul" in p.text and "XYZ projetés" not in p.text:
                p.text = (
                    "Entrées modèle : image RGB 224×224 (Zivid) + consigne texte ; jour 07 — variante YOLO "
                    "(label + XYZ dans le prompt) ; Grounding DINO (label seul ou label + XYZ projetés). "
                    "Jour 08 : boucle continue perception → action → réinférence."
                )
            break

    doc.save(DAY01)
    print(f"Updated {DAY01}")


if __name__ == "__main__":
    build_day08()
    update_day01_for_day08()

