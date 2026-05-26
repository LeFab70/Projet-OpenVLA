#!/usr/bin/env python3
"""Met à jour le texte du rapport final autour des figures YOLO / pipeline."""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DAY01 = "/Users/fabricekouonang/Documents/Projet-OpenVLA/OpenVLA_day01_stage_CCNB.docx"


def has_image(paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("a:blip")))


def clear_text_keep_images(paragraph) -> None:
    for run in list(paragraph.runs):
        if run._element.findall(".//" + qn("w:drawing")):
            continue
        run.text = ""


def insert_after(elem, doc, text, style=None):
    new_p = OxmlElement("w:p")
    elem.addnext(new_p)
    para = Paragraph(new_p, doc)
    if style:
        para.style = style
    para.add_run(text)
    return new_p


def main() -> None:
    doc = Document(DAY01)

    replacements = {
        "Figure 1 — Architecture conceptuelle OpenVLA (boucle VLA, entrées/sorties).": (
            "Figure 1 — Architecture conceptuelle OpenVLA (boucle VLA, entrées image + langage, "
            "sortie tokens d'action)."
        ),
        "Figure 2 — Répartition des poids par composant (VRAM minimale selon précision).": (
            "Figure 3 — Poids des modèles par composant (SigLIP, Prismatic, Llama 2 7B, action head)."
        ),
        "Figure 3 — Flux de données macro (instruction → capture → encodage → inférence → action → feedback).": (
            "Figure 4 — Vue macro du flux de données OpenVLA (instruction, capture, encodage SigLIP, "
            "décodeur Llama 2, action head, feedback)."
        ),
    }

    vram_text = (
        "Le décodeur LLM représente plus de 95 % du poids du modèle (~15 Go en FP16). "
        "Une GPU NVIDIA avec VRAM suffisante (≥ 16 Go) est requise pour l'inférence locale "
        "du modèle openvla-7b."
    )

    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            p.text = replacements[t]
        if has_image(p) and vram_text[:40] in p.text:
            clear_text_keep_images(p)

    # Légende Figure 2 (pipeline YOLO) avant image3 (p54)
    for i, p in enumerate(doc.paragraphs):
        if has_image(p):
            rels = [
                blip.get(qn("r:embed"))
                for blip in p._element.findall(".//" + qn("a:blip"))
            ]
            for rid in rels:
                target = doc.part.rels[rid].target_ref
                if target == "media/image3.png":
                    cap = "Figure 2 — Pipeline final Zivid + YOLO + OpenVLA (capture RGB/profondeur, détection 2D, coordonnées 3D, inférence, exécution UR)."
                    prev = p._element.getprevious()
                    if prev is not None and "Figure 2 — Pipeline" in (Paragraph(prev, p._parent).text or ""):
                        break
                    new_p = OxmlElement("w:p")
                    p._element.addprevious(new_p)
                    para = Paragraph(new_p, doc)
                    para.add_run(cap)
                    expl = (
                        "Étapes illustrées : (1) capture Zivid RGB + profondeur ; (2) détection 2D YOLOv8n "
                        "(ex. boule rouge) ; (3) pixel (u,v) du centre de la boîte ; (4) conversion Zivid "
                        "vers coordonnées 3D (X,Y,Z) en mètres ; (5–6) OpenVLA reçoit l'image 224×224 et une "
                        "consigne enrichie (label + position 3D) ; (7) le code RTDE exécute l'action prédite. "
                        "Scripts : scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py, returnAllPositions.py."
                    )
                    insert_after(new_p, doc, expl)
                    break

    # Texte explicatif après Figure 1 (image2)
    for p in doc.paragraphs:
        if has_image(p):
            for blip in p._element.findall(".//" + qn("a:blip")):
                if doc.part.rels[blip.get(qn("r:embed"))].target_ref == "media/image2.png":
                    note = (
                        "La figure 1 résume la boucle VLA : vision (image de la scène), langage (consigne "
                        "utilisateur) et action (tokens convertis en commandes robot). L'architecture interne "
                        "combine SigLIP/ViT, un projecteur MLP et Llama 2 7B."
                    )
                    nxt = p._element.getnext()
                    if nxt is None or note[:30] not in (Paragraph(nxt, p._parent).text or ""):
                        insert_after(p._element, doc, note)
                    insert_after(p._element, doc, vram_text)
                    break

    # II.1 — ajouter YOLO après paragraphe 224×224
    for i, p in enumerate(doc.paragraphs):
        if "Pour OpenVLA, l'image est convertie RGBA → RGB" in p.text:
            blocks = [
                (
                    "Heading 3",
                    "II.1.1 Détection d'objets YOLOv8 et coordonnées 3D (jour 07)",
                ),
                (
                    "Normal",
                    "Extension du démonstrateur : YOLOv8n (Ultralytics, yolov8n.pt) détecte les objets en 2D "
                    "sur l'image Zivid. Le centre de chaque boîte englobante (u, v) est projeté dans le nuage "
                    "de points pour obtenir (X, Y, Z) en mètres. L'objet le plus confiant alimente le prompt "
                    "OpenVLA (ex. pick up the ball at position X=… Y=… Z=…).",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/functions/returnAllPositions.py — capture + YOLO + liste 3D.",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py — pipeline complet Zivid → YOLO → OpenVLA.",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/test/test_yolo.py — test isolé YOLO sur image capturée.",
                ),
            ]
            anchor = p._element
            for style, text in reversed(blocks):
                anchor = insert_after(anchor, doc, text, style)
            break

    # II.2 — script YOLO
    for p in doc.paragraphs:
        if p.text.strip() == "Test intégré : scripts/integration/test/test_zivid_openvla.py.":
            insert_after(
                p._element,
                doc,
                "Pipeline YOLO + 3D : scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py.",
                "List Bullet",
            )
            break

    # Schéma cible
    for p in doc.paragraphs:
        if p.text.startswith("Schéma cible du démonstrateur"):
            p.text = (
                "Schéma cible du démonstrateur : Zivid (capture) → YOLOv8 (détection 2D) → coordonnées 3D "
                "→ prétraitement image + prompt enrichi → OpenVLA (inférence) → conversion action → UR "
                "(exécution RTDE) → nouvelle capture (boucle)."
            )
            break

    # II.5 — mention YOLO
    for p in doc.paragraphs:
        if p.text.startswith("Entrées modèle : image RGB"):
            p.text = (
                "Entrées modèle : image RGB 224×224 (Zivid) + consigne texte ; variante jour 07 avec label "
                "YOLO et position 3D injectés dans le prompt."
            )
            break

    # III.2 — YOLO
    for p in doc.paragraphs:
        if p.text.startswith("Stage : Zivid fournit 2D et 3D"):
            p.text = (
                "Stage : Zivid fournit 2D et 3D ; OpenVLA consomme le RGB 224×224. YOLOv8n relie la "
                "détection 2D au nuage de points pour cibler un objet (Figure 2). Le PLY reste disponible "
                "pour calibration et extensions."
            )
            break

    # TOC II.1.1
    for p in doc.paragraphs:
        if p.text.strip() == "   II.1  Perception : caméra Zivid MR130":
            nxt = p._element.getnext()
            if nxt is None or "II.1.1" not in (Paragraph(nxt, p._parent).text or ""):
                insert_after(
                    p._element,
                    doc,
                    "   II.1.1  Détection YOLOv8 + coordonnées 3D (jour 07)",
                )
            break

    # Pied de page
    for p in doc.paragraphs:
        if "Rapport final" in p.text and "Fabrice Kouonang" in p.text:
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — Parties I à III (II.5 + pipeline YOLO, jour 07)"
            )
            break

    doc.save(DAY01)
    print(f"Updated {DAY01}")


if __name__ == "__main__":
    main()
