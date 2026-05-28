#!/usr/bin/env python3
"""Reformate OpenVLA_day01_stage_CCNB.docx : page titre, TOC cliquable, IV/V."""

from __future__ import annotations

import copy
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FINAL = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")

INTRO_HEADING = "Rapport final de stage — Synthèse OpenVLA"
INTRO_TEXT = (
    "Ce document constitue le rapport final du stage réalisé au CCNB-INNOV. Il présente la veille "
    "technologique, l'expérimentation pratique et l'intégration d'un démonstrateur OpenVLA couplant un "
    "robot collaboratif UR et une caméra Zivid 2+ MR130. Les trois premières parties répondent aux "
    "objectifs initiaux du stage : architecture OpenVLA, composantes de la pipeline et comparaison aux "
    "approches robotiques traditionnelles."
)
CONTEXTE_HEADING = "Contexte et question directrice"
CONTEXTE_TEXT = (
    "Quels sont les ingrédients de la « sauce OpenVLA » et en quoi cette approche diffère-t-elle des "
    "méthodes robotiques plus traditionnelles ? Le stage combine veille technologique, expérimentation "
    "sur banc réel (Windows 11 Pro, GPU NVIDIA, Python 3.11) et documentation technique."
)
OBJECTIFS_HEADING = "Objectifs du stage"
OBJECTIFS = [
    "Étudier l'architecture générale d'OpenVLA et en produire une description technique claire et structurée.",
    "Identifier les différentes composantes de la pipeline : perception, inférence, génération d'actions et exécution robotique.",
    "Comparer l'approche OpenVLA aux approches robotiques traditionnelles (programmation, perception, planification, contrôle).",
    "Mettre en place un démonstrateur expérimental UR + Zivid.",
    "Réaliser l'intégration logicielle et matérielle du démonstrateur.",
    "Effectuer des essais, analyser les résultats et documenter limites et opportunités.",
    "Produire une documentation technique et une synthèse pour l'équipe CCNB-INNOV.",
    "Présenter les résultats sous forme de démonstration et présentation technique.",
]

TITLE_LINES = [
    ("Projet OpenVLA — Rapport final de stage — CCNB-INNOV", 14, True),
    ("Par : Fabrice Kouonang", 11, False),
    ("————————————————————————————", 10, False),
    ("Superviseur direct : Guillaume Batungwanayo", 11, False),
    ("Date : 14 mai 2026 | Rapport final", 10, False),
]


def _remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def _clone(paragraph):
    return copy.deepcopy(paragraph._element)


def _page_break_element():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _styled_paragraph_element(doc: Document, text: str, style: str):
    """Paragraphe avec style Word du modèle (ex. Titre1 = Heading 1)."""
    p = doc.add_paragraph(text, style=style)
    el = p._element
    doc.element.body.remove(el)
    return el


def _append_after(anchor, element):
    anchor.addnext(element)
    return element


def _add_toc_field(doc: Document, anchor_el):
    from docx.text.paragraph import Paragraph

    field_el = OxmlElement("w:p")
    anchor_el = _append_after(anchor_el, field_el)
    fp = Paragraph(field_el, doc.element.body)
    run = fp.add_run()
    for ftype in ("begin", "separate", "end"):
        if ftype == "separate":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = r'TOC \o "1-3" \h \z \u'
            run._r.append(instr)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), ftype)
        run._r.append(fld)
    return anchor_el


def _set_heading_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.text = text


def _find_paragraph(doc: Document, test_fn):
    for p in doc.paragraphs:
        if test_fn(p):
            return p
    return None


def _find_index(doc: Document, test_fn) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if test_fn(p):
            return i
    return None


def _extract_day11_indices(doc: Document) -> tuple[int, int] | None:
    start = _find_index(doc, lambda p: "II.6 Jour 11" in p.text)
    if start is None:
        return None
    end = start
    for i in range(start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if i > start and t.startswith("II.5") and "Flux" in t and doc.paragraphs[i].style.name == "Normal":
            end = i
            break
        if i > start and t.startswith("III. Comparaison") and doc.paragraphs[i].style.name == "Normal":
            end = i
            break
    else:
        end = len(doc.paragraphs)
    return start, end


def _extract_main_body_indices(doc: Document) -> int:
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("I. Architecture générale"):
            idx = i
    if idx is None:
        raise RuntimeError("Section I. Architecture introuvable.")
    return idx


def _clear_body_keep_logo(doc: Document):
    """Supprime tout le corps sauf le premier paragraphe (logo / en-tête)."""
    body = doc.element.body
    logo_el = doc.paragraphs[0]._element
    for child in list(body):
        if child is logo_el:
            continue
        body.remove(child)


def _build_front_matter(doc: Document, anchor_el):
    for text, size, bold in TITLE_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        el = p._element
        doc.element.body.remove(el)
        anchor_el = _append_after(anchor_el, el)

    anchor_el = _append_after(anchor_el, _page_break_element())

    h = _styled_paragraph_element(doc, "Table des matières", "Heading 1")
    anchor_el = _append_after(anchor_el, h)
    anchor_el = _add_toc_field(doc, anchor_el)
    anchor_el = _append_after(anchor_el, _page_break_element())

    for text, style in [
        (INTRO_HEADING, "Heading 1"),
        (INTRO_TEXT, None),
        (CONTEXTE_HEADING, "Heading 2"),
        (CONTEXTE_TEXT, None),
        (OBJECTIFS_HEADING, "Heading 2"),
    ]:
        p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
        el = p._element
        doc.element.body.remove(el)
        anchor_el = _append_after(anchor_el, el)

    for obj in OBJECTIFS:
        p = doc.add_paragraph(obj, style="List Bullet")
        el = p._element
        doc.element.body.remove(el)
        anchor_el = _append_after(anchor_el, el)

    anchor_el = _append_after(anchor_el, _page_break_element())
    return anchor_el


def _move_iv_v_after_part_iii(doc: Document):
    """Place IV/V après la partie III (contenu en fin de rapport)."""
    h_iii = _find_paragraph(
        doc, lambda p: p.style.name == "Heading 1" and p.text.strip().startswith("III. Comparaison")
    )
    if not h_iii:
        return

    h_iv_start = _find_paragraph(doc, lambda p: p.text.strip() == "IV. Tests et intégration")
    if not h_iv_start:
        h_iv_start = _find_paragraph(doc, lambda p: p.text.strip().startswith("IV.1 Flux"))
    if not h_iv_start:
        return

    h_iii_el = h_iii._element
    h_iv_el = h_iv_start._element

    h_iii_end = None
    past_iii = False
    for p in doc.paragraphs:
        if p._element is h_iii_el:
            past_iii = True
            continue
        if not past_iii:
            continue
        if p.text.strip().startswith("Les parties IV") or p.text.strip().startswith("Programmeur :"):
            break
        h_iii_end = p
    if not h_iii_end:
        return

    body = doc.element.body
    moving = []
    capture = False
    for p in doc.paragraphs:
        if p._element is h_iii_el:
            break
        if p._element is h_iv_el:
            capture = True
        if capture:
            moving.append(p._element)

    for el in moving:
        body.remove(el)

    prev = h_iii_end._element
    for el in moving:
        prev.addnext(el)
        prev = el


def _rename_day11_subsections(doc: Document):
    sub = 0
    in_v2 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style.name == "Heading 2" and t.startswith("V.2 Jour 11"):
            in_v2 = True
            continue
        if not in_v2:
            continue
        if p.style.name == "Heading 1" or (p.style.name == "Heading 2" and not t.startswith("V.2.")):
            break
        if p.style.name == "Heading 3":
            sub += 1
            label = t.split(".", 1)[-1].strip() if "." in t[:5] else t
            _set_heading_text(p, f"V.2.{sub} {label}")


def _restructure_iv_v(doc: Document, day11_elements: list):
    for p in doc.paragraphs:
        t = p.text.strip()
        if "II.5" in t and "Flux de données" in t:
            _set_heading_text(p, "IV.1 Flux de données robot, Zivid et interprétation OpenVLA (jour 07)")
            p.style = doc.styles["Heading 2"]

    mapping = [
        ("Jour 08 — boucle continue", "IV.2 Jour 08 — Boucle continue perception → action → réinférence", 2),
        ("Jour 09 :", "IV.3 Jour 09 — Interprétation des résultats de la boucle continue", 2),
        ("Jour 10 — refactoring", "V.1 Jour 10 — Refactoring pipeline (modules pipeline/)", 2),
    ]
    for prefix, new_title, level in mapping:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                _set_heading_text(p, new_title)
                p.style = doc.styles[f"Heading {level}"]
                break

    h_iv1 = _find_paragraph(doc, lambda p: p.text.strip().startswith("IV.1 Flux"))
    if h_iv1 and not _find_paragraph(doc, lambda p: p.text.strip() == "IV. Tests et intégration"):
        h_iv1._element.addprevious(_styled_paragraph_element(doc, "IV. Tests et intégration", "Heading 1"))

    h_v1 = _find_paragraph(doc, lambda p: p.text.strip().startswith("V.1 Jour 10"))
    if h_v1 and not _find_paragraph(doc, lambda p: p.text.strip() == "V. Pipeline final du démonstrateur"):
        h_v1._element.addprevious(_styled_paragraph_element(doc, "V. Pipeline final du démonstrateur", "Heading 1"))

    if not day11_elements:
        return

    h_v1 = _find_paragraph(doc, lambda p: p.text.strip().startswith("V.1 Jour 10"))
    if not h_v1:
        return
    anchor = h_v1._element
    for p in doc.paragraphs:
        if p.text.strip().startswith("Points critiques"):
            anchor = p._element
            break

    v2_h = _styled_paragraph_element(doc, "V.2 Jour 11 — Calibration caméra → robot (T_tcp_cam)", "Heading 2")
    anchor.addnext(v2_h)
    prev = v2_h
    for el in day11_elements:
        text = "".join(el.itertext()).strip()
        if text.startswith("II.6 Jour 11"):
            continue
        prev.addnext(el)
        prev = el
def _update_footer(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Les parties IV à VII"):
            p.text = (
                "Les parties IV (tests et intégration) et V (pipeline final du démonstrateur, jours 08 à 11) "
                "complètent ce rapport. Les sections VI–VII (documentation et présentation) seront finalisées "
                "en fin de stage."
            )
        if t.startswith("Programmeur : Fabrice Kouonang"):
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — parties I à V"
            )


def reformat():
    src = Document(FINAL)

    day11_range = _extract_day11_indices(src)
    day11_elements = []
    if day11_range:
        s, e = day11_range
        day11_elements = [_clone(src.paragraphs[i]) for i in range(s, e)]

    body_start = _extract_main_body_indices(src)
    body_elements = [_clone(src.paragraphs[i]) for i in range(body_start, len(src.paragraphs))]

    new_doc = Document(FINAL)
    _clear_body_keep_logo(new_doc)

    anchor = new_doc.paragraphs[0]._element
    anchor = _build_front_matter(new_doc, anchor)
    for el in body_elements:
        anchor = _append_after(anchor, el)

    _restructure_iv_v(new_doc, day11_elements)
    _move_iv_v_after_part_iii(new_doc)
    _rename_day11_subsections(new_doc)
    _update_footer(new_doc)

    new_doc.save(FINAL)
    print(f"Reformatted {FINAL}")
    print("Word : clic droit sur la table des matières → Mettre à jour les champs → Table entière.")


if __name__ == "__main__":
    reformat()
