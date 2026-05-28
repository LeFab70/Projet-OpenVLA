#!/usr/bin/env python3
"""Génère OpenVLA_day11_rapport_hebdomadaire.docx (rapport hebdomadaire)."""

import copy
import os
import shutil

from docx import Document

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(BASE, "OpenVLA_day05_openvla_integration.docx")
DAY11 = os.path.join(BASE, "OpenVLA_day11_rapport_hebdomadaire.docx")


def clear_body_keep_logo(doc: Document) -> None:
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def build():
    shutil.copy2(TEMPLATE, DAY11)
    doc = Document(DAY11)
    clear_body_keep_logo(doc)

    add_para(doc, "Projet OpenVLA — Jour 11 — Rapport hebdomadaire (semaine 26–29 mai 2026)")
    add_para(doc, "Par : Fabrice Kouonang")
    add_para(doc, "————————————————————————————")
    add_para(doc, "Superviseur direct : Guillaume Batungwanayo")
    add_para(doc, "Date : 29 mai 2026 | Jour : 11")

    add_heading(doc, "1. Résumé de la semaine (ce qui a été fait)")
    add_bullet(doc, "Jour 08 : boucle continue perception → action → réinférence (OpenVLA adaptatif).")
    add_bullet(doc, "Jour 09 : interprétation des logs (demo_adaptatif_openvla_print_value.py).")
    add_bullet(doc, "Jour 10 : refactoring en modules indépendants (dossier pipeline/).")
    add_bullet(doc, "Correction bug DINO : mapping pixels image → point cloud (résolutions 2D/3D Zivid différentes).")
    add_bullet(doc, "Création du script de calibration main-œil : pipeline/calibrer_robot.py → fichier T_tcp_cam.npy.")

    add_heading(doc, "2. Problèmes rencontrés")
    add_bullet(doc, "Calibration main-œil absente : sans T_tcp_cam.npy, la caméra et le robot ne partagent pas le même repère.")
    add_bullet(doc, "IndexError dans dino_detector.py : coordonnées pixel (u,v) hors limites du nuage de points.")
    add_bullet(doc, "Mode réel bloqué tant que la matrice de calibration n’est pas générée et validée.")

    add_heading(doc, "3. Solutions proposées / mises en œuvre")
    add_heading(doc, "3.0 Architecture calibration — I, II, III", level=2)
    add_para(
        doc,
        "Deux modules distincts dans pipeline/ : un script autonome pour créer T_tcp_cam.npy, "
        "et un module chargé à chaque exécution du pipeline pour convertir les coordonnées pendant les déplacements OpenVLA.",
    )
    add_heading(doc, "I. Fichier T_tcp_cam.npy (traducteur caméra → robot)", level=3)
    add_para(
        doc,
        "Matrice homogène 4×4 sauvegardée avec numpy (np.save) après calibration Zivid réussie. "
        "Sans ce fichier, DINO/OpenVLA voient l’objet dans le repère caméra ; le UR16e ne peut pas "
        "interpréter « à gauche de la caméra » sans conversion.",
    )
    add_bullet(doc, "Chemin : CALIBRATION_FILE dans pipeline/config.py → T_tcp_cam.npy.")
    add_heading(doc, "II. pipeline/calibrer_robot.py — exécution seule (création du fichier)", level=3)
    add_para(
        doc,
        "Script indépendant de main_real.py / main_sim.py. À lancer une fois (ou après déplacement de la caméra) "
        "pour produire T_tcp_cam.npy.",
    )
    add_bullet(
        doc,
        "Pour chaque pose : capture Zivid (camera.capture()), détection mire damier noir/blanc "
        "(zivid.calibration.detect_feature_points sur le point cloud).",
    )
    add_bullet(
        doc,
        "Échantillons hand-eye : pose TCP UR + détection valide → calibrate_eye_in_hand() → np.save(T_tcp_cam).",
    )
    add_bullet(doc, "Commande : python -m pipeline.calibrer_robot (environnements Zivid + RTDE).")
    add_heading(doc, "III. pipeline/calibration.py — à chaque run du pipeline OpenVLA", level=3)
    add_bullet(doc, "load_calibration() : charge T_tcp_cam.npy (ou matrice défaut si absent).")
    add_bullet(
        doc,
        "cam_to_robot() : reconvertit xyz caméra → base robot à chaque déplacement (utilisé par main_real.py).",
    )
    add_bullet(doc, "compute_distance_tcp_to_object() : distance TCP–objet pour critère d’arrêt / prompt.")
    add_para(
        doc,
        "calibrer_robot.py trouve la mire et écrit le fichier ; calibration.py le relit en continu pendant la boucle OpenVLA.",
        bold=True,
    )

    add_heading(doc, "3.1 Le fichier .npy — traducteur mathématique (T_tcp_cam)", level=2)
    add_para(
        doc,
        "Le fichier .npy contient la matrice de transformation homogène T_tcp_cam (4×4). "
        "Sans lui, la caméra et le bras « parlent deux langues différentes » :",
    )
    add_bullet(doc, "La caméra dit : « Je vois la bouteille à 10 cm à ma gauche et 50 cm devant moi. »")
    add_bullet(doc, "Le robot dit : « Moi, je suis à 40 cm de ma base. Je ne sais pas où est ta gauche. »")
    add_para(
        doc,
        "La matrice T_tcp_cam permet de calculer la correspondance : "
        "« la gauche de la caméra correspond en fait à l’axe -Y de la base du robot ».",
    )
    add_para(doc, "Fichier cible : CALIBRATION_FILE (config.py) → T_tcp_cam.npy", bold=True)

    add_heading(doc, "3.2 Script de calibration — pipeline/calibrer_robot.py", level=2)
    add_bullet(doc, "Capture de poses variées avec mire Zivid (hand-eye, eye-in-hand).")
    add_bullet(doc, "Calcul et sauvegarde de T_tcp_cam via zivid.calibration.calibrate_eye_in_hand().")
    add_bullet(doc, "Utilisation ensuite par pipeline/calibration.py et main_real.py.")

    add_heading(doc, "3.3 Les 3 règles d’or pour les poses de calibration", level=2)
    add_bullet(
        doc,
        "Varier l’inclinaison : ne pas rester toujours à la verticale (RX=0, RY=3.14). "
        "Incliner le poignet de ±15° à ±25° (environ 0,3 à 0,4 rad).",
    )
    add_bullet(doc, "Varier la hauteur (Z) : photos à environ 40 cm, 50 cm et 60 cm de la mire.")
    add_bullet(doc, "Rotation autour de l’axe Z (RZ) : faire tourner l’outil sur lui-même.")

    add_heading(doc, "4. Objectifs semaine suivante")
    add_bullet(doc, "Exécuter la calibration sur le UR16e et valider T_tcp_cam.npy en conditions réelles.")
    add_bullet(doc, "Tester pipeline/main_real.py avec calibration réelle + workspace.")
    add_bullet(doc, "Poursuivre les rapports hebdomadaires dans le dossier rapports_hebdomadaires/ (détails à compléter).")
    add_bullet(doc, "Valider la prise d’objet en boucle fermée (DINO + OpenVLA + Robotiq).")

    add_heading(doc, "Livrables jour 11")
    add_bullet(doc, "OpenVLA_day11_rapport_hebdomadaire.docx (ce document).")
    add_bullet(doc, "rapports_hebdomadaires/ — dossier pour les prochains rapports.")
    add_bullet(doc, "pipeline/calibrer_robot.py — génération T_tcp_cam.npy.")
    add_bullet(doc, "README — section Jour 11 mise à jour.")

    doc.save(DAY11)
    print(f"Built {DAY11}")


if __name__ == "__main__":
    build()
