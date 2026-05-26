#!/usr/bin/env python3
"""Ajoute la comparaison OpenVLA / robotique classique (notes captures) aux rapports."""

import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DAY01 = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")
DAY02 = os.path.join(BASE, "OpenVLA_day02_prise_en_main.docx")
DAY07 = os.path.join(BASE, "OpenVLA_day07_robot_data.docx")

MARKER_III1 = "III.1 — Robotique classique seule"
MARKER_DAY02 = "OpenVLA vs robotique classique (synthèse)"
MARKER_DAY07 = "9. OpenVLA vs robotique classique"


def insert_after(elem, doc, text, style=None):
    new_p = OxmlElement("w:p")
    elem.addnext(new_p)
    para = Paragraph(new_p, doc)
    if style:
        para.style = style
    para.add_run(text)
    return new_p


def has_marker(doc, marker: str) -> bool:
    return any(marker in p.text for p in doc.paragraphs)


def remove_paragraphs_containing(doc, needles):
    for p in list(doc.paragraphs):
        if any(n in p.text for n in needles):
            p._element.getparent().remove(p._element)


def update_day01() -> None:
    doc = Document(DAY01)
    remove_paragraphs_containing(
        doc,
        [
            MARKER_III1,
            "III.1 — Ce qu'OpenVLA ajoute",
            "OpenVLA peut aider le robot à :",
            "En robotique classique seule, le robot peut aller précisément",
            "Compréhension du langage (consignes en langage naturel)",
            "Stratégie et planification implicite dans le modèle VLA",
            "Adaptation et comportement contextuel",
            "Exemples de consignes que le classique peine",
            "Le démonstrateur du stage reste hybride : RTDE/URScript",
            "Robot classique — il sait où est la cible",
            "OpenVLA peut comprendre une consigne du type",
            "OpenVLA aide surtout à :",
            "Raisonner et planifier à partir de la scène",
            "Adapter le mouvement et corriger en boucle fermée",
            "Gérer des scènes complexes et comprendre le langage humain",
            "Tandis que Zivid + nuage de points",
        ],
    )

    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "III.1" in p.text:
            anchor = p._element
            blocks = [
                (
                    "Normal",
                    "OpenVLA peut aider le robot à : mieux comprendre la scène ; mieux choisir "
                    "l'action ; mieux adapter le mouvement.",
                ),
                ("Heading 3", MARKER_III1),
                (
                    "Normal",
                    "En robotique classique seule, le robot peut aller précisément au XYZ. "
                    "Mais il ne comprend pas forcément quoi prendre, dans quel ordre, ni comment "
                    "s'adapter au contexte.",
                ),
                ("Heading 3", "III.1 — Ce qu'OpenVLA ajoute"),
                ("List Bullet", "Compréhension du langage (consignes en langage naturel)."),
                ("List Bullet", "Stratégie et planification implicite dans le modèle VLA."),
                ("List Bullet", "Adaptation et comportement contextuel."),
                (
                    "Normal",
                    "Exemples de consignes que le classique peine à interpréter seul : "
                    "« la tasse est derrière un obstacle » ; « ouvrir le tiroir avant ».",
                ),
                (
                    "Normal",
                    "Le démonstrateur du stage reste hybride : RTDE/URScript pour l'exécution "
                    "prévisible, OpenVLA pour la décision image + langage → action.",
                ),
            ]
            for style, text in blocks:
                anchor = insert_after(anchor, doc, text, style)
            break

    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "III.2" in p.text:
            anchor = p._element
            blocks = [
                (
                    "Normal",
                    "Robot classique — il sait où est la cible (ex. où est la balle), mais il peut "
                    "heurter l'obstacle, ne pas comprendre la scène dans son ensemble, ou échouer "
                    "si la trajectoire change.",
                ),
                (
                    "Normal",
                    "OpenVLA peut comprendre une consigne du type « il y a un obstacle » et orienter "
                    "la décision : contourner, déplacer un objet, changer l'approche, réessayer "
                    "autrement.",
                ),
            ]
            for style, text in blocks:
                anchor = insert_after(anchor, doc, text, style)
            break

    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "III.3" in p.text:
            anchor = insert_after(p._element, doc, "OpenVLA aide surtout à :", "Normal")
            for text in [
                "Raisonner et planifier à partir de la scène et du langage.",
                "Adapter le mouvement et corriger en boucle fermée.",
                "Gérer des scènes complexes et comprendre le langage humain.",
            ]:
                anchor = insert_after(anchor, doc, text, "List Bullet")
            break

    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "III.4" in p.text:
            nxt = p._element.getnext()
            if nxt is not None and "Tandis que Zivid" in (Paragraph(nxt, p._parent).text or ""):
                break
            insert_after(
                p._element,
                doc,
                "Tandis que Zivid + nuage de points fournissent la géométrie précise (RGB, "
                "profondeur, XYZ en mètres), OpenVLA apporte la couche sémantique et décisionnelle "
                "au-dessus de cette perception.",
                "Normal",
            )
            break

    for p in doc.paragraphs:
        if "Rapport final" in p.text and "Fabrice Kouonang" in p.text:
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — I à III (comparaison OpenVLA / robotique classique)"
            )
            break

    doc.save(DAY01)
    print(f"Updated {DAY01}")


def update_day02() -> None:
    doc = Document(DAY02)
    if has_marker(doc, MARKER_DAY02):
        print(f"Skip {DAY02} (déjà à jour)")
        return

    for p in doc.paragraphs:
        if p.text.startswith("OpenVLA") and "pont entre la vision" in p.text:
            anchor = p._element
            blocks = [
                ("Heading 1", MARKER_DAY02),
                (
                    "Normal",
                    "Robotique classique : le robot atteint précisément un XYZ, mais ne décide pas "
                    "seul quoi prendre, dans quel ordre, ni comment s'adapter. OpenVLA ajoute la "
                    "compréhension du langage, la stratégie, l'adaptation et un comportement "
                    "contextuel (ex. obstacle, tiroir à ouvrir avant).",
                ),
                (
                    "Normal",
                    "OpenVLA peut aider le robot à mieux comprendre la scène, mieux choisir "
                    "l'action et mieux adapter le mouvement — complément de la géométrie Zivid.",
                ),
            ]
            for style, text in blocks:
                anchor = insert_after(anchor, doc, text, style)
            break

    doc.save(DAY02)
    print(f"Updated {DAY02}")


def update_day07() -> None:
    doc = Document(DAY07)
    if has_marker(doc, MARKER_DAY07):
        print(f"Skip {DAY07} (déjà à jour)")
        return

    for p in doc.paragraphs:
        if p.text.strip() == "6. Synthèse — ce qu'OpenVLA interprète":
            anchor = p._element
            blocks = [
                ("Heading 1", MARKER_DAY07),
                (
                    "Normal",
                    "Rappel pédagogique (comparaison au classique) : la Zivid et le nuage de points "
                    "donnent où sont les objets en 3D ; OpenVLA interprète la scène et la consigne "
                    "pour choisir une action (image 224×224 + prompt, pas le PLY brut).",
                ),
                (
                    "Normal",
                    "OpenVLA peut aider le robot à : mieux comprendre la scène ; mieux choisir "
                    "l'action ; mieux adapter le mouvement — par exemple face à un obstacle ou une "
                    "consigne du type « ouvrir le tiroir avant ».",
                ),
            ]
            for style, text in blocks:
                anchor = insert_after(anchor, doc, text, style)
            break

    doc.save(DAY07)
    print(f"Updated {DAY07}")


def main() -> None:
    update_day01()
    update_day02()
    update_day07()


if __name__ == "__main__":
    main()
