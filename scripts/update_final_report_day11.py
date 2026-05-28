#!/usr/bin/env python3
"""Ajoute la section Jour 11 (calibration) au rapport final OpenVLA_day01_stage_CCNB.docx."""

import os

from docx import Document
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FINAL = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")
MARKER = "II.6 Jour 11 — Calibration caméra → robot (T_tcp_cam)"


def _insert_after(paragraph, text: str, style=None, bold=False):
    new_p = paragraph.insert_paragraph_before(text, style=style)
    if bold and new_p.runs:
        new_p.runs[0].bold = True
    return new_p


def _add_block(anchor, blocks):
    """Insert blocks before anchor (reverse order so first block ends up right after anchor)."""
    current = anchor
    for kind, content in reversed(blocks):
        if kind == "h2":
            current = _insert_after(current, content, style="Heading 2")
        elif kind == "h3":
            current = _insert_after(current, content, style="Heading 3")
        elif kind == "bullet":
            current = _insert_after(current, content, style="List Bullet")
        elif kind == "para":
            current = _insert_after(current, content)
        elif kind == "parab":
            current = _insert_after(current, content, bold=True)
    return current


def already_has_day11(doc: Document) -> bool:
    return any(MARKER in p.text for p in doc.paragraphs)


def build_blocks():
    return [
        ("h2", MARKER),
        (
            "para",
            "Le 29 mai 2026 (jour 11), le démonstrateur intègre la calibration main-œil (eye-in-hand) "
            "UR16e + Zivid. Cette section distingue la création du fichier de calibration (script autonome) "
            "de son utilisation à chaque exécution du pipeline OpenVLA.",
        ),
        ("h3", "I. Le fichier T_tcp_cam.npy — traducteur entre caméra et robot"),
        (
            "para",
            "La matrice homogène 4×4 T_tcp_cam relie le repère de la caméra Zivid (montée sur le TCP) "
            "au repère outil du robot. Sans ce fichier, les coordonnées 3D issues de Grounding DINO "
            "(repère caméra) ne peuvent pas être converties en positions fiables pour le UR16e.",
        ),
        ("bullet", "Caméra : « l’objet est à gauche et devant moi » (repère optique Zivid)."),
        ("bullet", "Robot : « je me déplace dans la base UR » — il ne connaît pas la « gauche » de la caméra."),
        (
            "para",
            "Fichier produit : T_tcp_cam.npy (chemin CALIBRATION_FILE dans pipeline/config.py). "
            "Sauvegarde avec numpy : np.save(CALIBRATION_FILE, T_tcp_cam) après calibration réussie.",
        ),
        ("h3", "II. pipeline/calibrer_robot.py — script autonome (création du fichier)"),
        (
            "para",
            "Ce script s’exécute seul, indépendamment de main_real.py et main_sim.py. "
            "Il sert uniquement à générer ou régénérer T_tcp_cam.npy avant les runs du pipeline.",
        ),
        ("bullet", "Connexion Zivid (zivid.Application, connect_camera) et UR16e (RTDE)."),
        (
            "bullet",
            "Pour chaque pose de calibration : déplacement moveL, lecture TCP, capture Zivid "
            "(camera.capture()), détection de la mire damier noir et blanc via "
            "zivid.calibration.detect_feature_points(frame.point_cloud()).",
        ),
        (
            "bullet",
            "Accumulation des échantillons HandEyeInput (pose robot + détection valide), puis "
            "zivid.calibration.calibrate_eye_in_hand(detection_results).",
        ),
        (
            "bullet",
            "Si result.valid() : extraction de T_tcp_cam = result.transform() et np.save vers T_tcp_cam.npy.",
        ),
        (
            "para",
            "Lancement typique (depuis la racine du projet, environnement Zivid + RTDE) : "
            "python -m pipeline.calibrer_robot",
        ),
        ("h3", "III. pipeline/calibration.py — utilisé à chaque run OpenVLA"),
        (
            "para",
            "À chaque exécution du pipeline (main_real.py, boucle adaptative), calibration.py charge "
            "T_tcp_cam.npy via load_calibration() et convertit les coordonnées caméra → base robot "
            "avec cam_to_robot(xyz_cam_m, tcp_pose, T_tcp_cam) lors des déplacements guidés par OpenVLA.",
        ),
        ("bullet", "load_calibration() : np.load si le fichier existe ; sinon matrice identité (mode non calibré)."),
        (
            "bullet",
            "cam_to_robot() : chaîne T_base_tcp @ T_tcp_cam @ P_cam pour placer l’objet détecté dans la base UR.",
        ),
        (
            "bullet",
            "compute_distance_tcp_to_object() : distance restante TCP → objet (critère d’arrêt / prompt adaptatif).",
        ),
        (
            "para",
            "Distinction essentielle : calibrer_robot.py crée le fichier une fois (ou après déplacement de la caméra) ; "
            "calibration.py réutilise ce fichier à chaque cycle perception → OpenVLA → RTDE.",
        ),
        ("h3", "IV. Les 3 règles d’or pour les poses de calibration"),
        (
            "bullet",
            "Varier l’inclinaison : ne pas rester toujours à la verticale (RX=0, RY=3.14) ; incliner ±15° à ±25°.",
        ),
        ("bullet", "Varier la hauteur Z : captures à environ 40 cm, 50 cm et 60 cm de la mire."),
        ("bullet", "Rotation RZ : faire tourner l’outil sur lui-même pour enrichir les échantillons hand-eye."),
        ("h3", "V. Rapport hebdomadaire et livrables jour 11"),
        (
            "para",
            "Rapport détaillé : OpenVLA_day11_rapport_hebdomadaire.docx. "
            "Dossier rapports_hebdomadaires/ pour les synthèses de fin de semaine.",
        ),
        ("bullet", "Semaine 26–29 mai : boucle continue (jour 08), interprétation logs (jour 09), refactoring pipeline (jour 10)."),
        ("bullet", "Correction DINO : mapping pixel image → résolution point cloud Zivid (scale_u, scale_v)."),
        ("bullet", "Objectif semaine suivante : valider T_tcp_cam.npy en réel et tester main_real.py avec workspace."),
    ]


def update():
    doc = Document(FINAL)
    if already_has_day11(doc):
        print(f"Section déjà présente dans {FINAL}, rien à ajouter.")
        return

    anchor = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("III. Comparaison OpenVLA"):
            anchor = doc.paragraphs[i - 1] if i > 0 else p
            break
    if anchor is None:
        for p in reversed(doc.paragraphs):
            if "Jour 08" in p.text or "demo_adaptatif" in p.text:
                anchor = p
                break
    if anchor is None:
        raise RuntimeError("Point d'insertion introuvable dans le rapport final.")

    _add_block(anchor, build_blocks())

    for p in doc.paragraphs:
        if p.text.strip().startswith("Programmeur : Fabrice Kouonang") and "I à III" in p.text:
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — I à III + jour 11 (calibration T_tcp_cam)"
            )
            break

    doc.save(FINAL)
    print(f"Updated {FINAL}")


if __name__ == "__main__":
    update()
