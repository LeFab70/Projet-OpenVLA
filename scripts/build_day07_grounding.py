#!/usr/bin/env python3
"""Reconstruit OpenVLA_day07 et met à jour le rapport final (YOLO vs Grounding DINO)."""

import copy
import os
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(BASE, "OpenVLA_day05_openvla_integration.docx")
DAY07 = os.path.join(BASE, "OpenVLA_day07_robot_data.docx")
DAY01 = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")


def logo_paragraph_element():
    doc5 = Document(TEMPLATE)
    return copy.deepcopy(doc5.paragraphs[0]._element)


def clear_body_keep_logo(doc: Document) -> None:
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def insert_after(elem, doc, text, style=None):
    new_p = OxmlElement("w:p")
    elem.addnext(new_p)
    para = Paragraph(new_p, doc)
    if style:
        para.style = style
    para.add_run(text)
    return new_p


def remove_paragraphs_containing(doc, needles):
    for p in list(doc.paragraphs):
        if any(n in p.text for n in needles):
            p._element.getparent().remove(p._element)


def build_day07():
    shutil.copy2(TEMPLATE, DAY07)
    doc = Document(DAY07)
    clear_body_keep_logo(doc)

    add_para(doc, "Projet OpenVLA — Jour 07 — Données robot, YOLO, Grounding DINO et OpenVLA")
    add_para(doc, "Par : Fabrice Kouonang")
    add_para(doc, "————————————————————————————")
    add_para(doc, "Superviseur direct : Guillaume Batungwanayo")
    add_para(doc, "Date : 25 mai 2026 | Jour : 07")

    add_heading(doc, "Objectif de la journée")
    add_para(
        doc,
        "Comprendre le flux de données entre Zivid, OpenVLA et le robot UR ; comparer deux "
        "approches de détection 2D (YOLOv8n et Grounding DINO) pour extraire des coordonnées 3D "
        "et alimenter OpenVLA ; clarifier ce qui entre dans le modèle versus ce que le contrôleur "
        "robot conserve.",
    )

    add_heading(doc, "Contexte")
    add_para(
        doc,
        "Le jour 06 a validé la boucle Zivid → OpenVLA → UR16e (demoTest.py). Le jour 07 approfondit "
        "les entrées du modèle et ajoute une couche de perception intermédiaire : détection d'objets "
        "en 2D puis projection 3D via le nuage de points Zivid. Deux détecteurs ont été testés : "
        "YOLOv8n (classes fixes COCO) et Grounding DINO (vocabulaire ouvert par texte).",
    )

    add_heading(doc, "1. Flux de données — rappel (OpenVLA et robot UR)")
    add_bullet(doc, "Entrées OpenVLA : image RGB 224×224 (Zivid) + consigne texte (prompt VLA).")
    add_bullet(doc, "Sortie OpenVLA : 7D (dx, dy, dz, rx, ry, rz, gripper).")
    add_bullet(doc, "Robot UR (RTDE) : pose TCP lue après inférence — non envoyée au réseau.")
    add_bullet(doc, "Formule d'exécution : new_pose = current_pose + action × SCALE (SCALE = 0,05).")
    add_para(
        doc,
        "Prompt type : In: What action should the robot take to {INSTRUCTION}?\\nOut:",
        bold=True,
    )

    add_heading(doc, "2. Pipeline commun Zivid → détection 2D → 3D")
    add_para(
        doc,
        "Les deux pipelines partagent les mêmes étapes initiales :",
    )
    add_bullet(doc, "Capture Zivid 2D/3D (RGB + nuage de points xyz).")
    add_bullet(doc, "Détection 2D : boîte englobante → centre pixel (u, v).")
    add_bullet(doc, "Projection : (u, v) mappé vers le point cloud → (X, Y, Z) en mètres.")
    add_bullet(doc, "Sélection de l'objet le plus confiant pour la suite du pipeline.")
    add_bullet(doc, "OpenVLA : image 224×224 + prompt → predict_action (bridge_orig).")

    add_heading(doc, "3. YOLOv8n — détection à classes fixes")
    add_heading(doc, "3.1 Principe", level=2)
    add_para(
        doc,
        "YOLOv8n (Ultralytics, yolov8n.pt) détecte des objets parmi les classes COCO prédéfinies "
        "(bottle, cup, cell phone, etc.). Le seuil de confiance utilisé : YOLO_CONF = 0,25.",
    )
    add_heading(doc, "3.2 Prompt OpenVLA (YOLO)", level=2)
    add_para(
        doc,
        "Les coordonnées 3D sont injectées dans la consigne texte envoyée à OpenVLA :",
    )
    add_para(
        doc,
        "pick up the {label} at position X=…m Y=…m Z=…m",
        bold=True,
    )
    add_bullet(doc, "Script : scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py")
    add_bullet(doc, "Fonction réutilisable : scripts/openVLA_ZIVID/functions/returnAllPositions.py")
    add_bullet(doc, "Test isolé : scripts/openVLA_ZIVID/test/test_yolo.py")

    add_heading(doc, "4. Grounding DINO — détection open-vocabulary")
    add_heading(doc, "4.1 Principe", level=2)
    add_para(
        doc,
        "Grounding DINO (IDEA-Research/grounding-dino-base) associe image et texte libre. "
        "L'utilisateur fournit une requête textuelle (ex. « cell phone. ») ; le modèle retourne "
        "les boîtes correspondantes sans liste de classes fixe. Seuils : BOX_THRESHOLD = 0,35, "
        "TEXT_THRESHOLD = 0,25.",
    )
    add_heading(doc, "4.2 Prompt OpenVLA (Grounding DINO)", level=2)
    add_para(
        doc,
        "Choix de conception du test : OpenVLA reçoit uniquement la description sémantique de "
        "l'objet — les coordonnées 3D ne sont pas injectées dans le prompt. Elles restent "
        "disponibles pour le contrôleur robot (affichage / exécution future).",
    )
    add_para(doc, "pick up the {label}", bold=True)
    add_bullet(doc, "Script : scripts/openVLA_ZIVID/test/test_zivid_groundingDino.py")

    add_heading(doc, "5. Comparaison YOLO vs Grounding DINO")
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    rows = [
        ("Critère", "YOLOv8n", "Grounding DINO"),
        ("Type de détection", "Classes COCO fixes", "Open-vocabulary (texte libre)"),
        ("Entrée détecteur", "Image RGB", "Image RGB + prompt texte (ex. cell phone.)"),
        ("Modèle", "yolov8n.pt (Ultralytics)", "grounding-dino-base (Hugging Face)"),
        ("Prompt OpenVLA", "Label + position 3D (X, Y, Z)", "Label seul (sans XYZ)"),
        ("Coordonnées 3D", "Dans le prompt VLA", "Réservées au contrôleur robot"),
        ("Cas d'usage stage", "Objets COCO sur table de démo", "Ciblage par consigne naturelle"),
    ]
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_heading(doc, "6. Synthèse — ce qu'OpenVLA interprète")
    add_para(
        doc,
        "Dans les deux cas, OpenVLA ne reçoit pas directement le nuage de points ni la pose UR : "
        "il voit l'image 224×224 et le texte du prompt. La différence jour 07 est la richesse "
        "du prompt (avec ou sans XYZ) et le mode de détection amont (classes fixes vs texte libre). "
        "La géométrie 3D sert à localiser l'objet ; la décision motrice reste un delta 7D appris.",
    )

    add_heading(doc, "7. Scripts de référence")
    add_bullet(doc, "scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py — Zivid → YOLO → OpenVLA (XYZ dans prompt).")
    add_bullet(doc, "scripts/openVLA_ZIVID/test/test_zivid_groundingDino.py — Zivid → Grounding DINO → OpenVLA.")
    add_bullet(doc, "scripts/integration/testUR_ZIVID/demoTest.py — boucle UR sans couche détection.")
    add_bullet(doc, "scripts/integration/test/test_zivid_openvla.py — Zivid + OpenVLA sans détecteur.")

    add_heading(doc, "8. Pistes et limites")
    add_bullet(doc, "YOLO : rapide, mais limité aux classes COCO ; mauvais label si objet hors vocabulaire.")
    add_bullet(doc, "Grounding DINO : flexible (texte libre), plus lourd ; nécessite GPU et transformers.")
    add_bullet(doc, "Fusion future : Grounding DINO pour cibler + YOLO pour vitesse, ou XYZ toujours côté RTDE.")

    add_heading(doc, "Livrables jour 07")
    add_bullet(doc, "OpenVLA_day07_robot_data.docx (ce document).")
    add_bullet(doc, "Rapport final : sections II.1.1 (YOLO) et II.1.2 (Grounding DINO).")
    add_bullet(doc, "README — entrée Jour 07 mise à jour.")

    doc.save(DAY07)
    print(f"Built {DAY07}")


def update_day01():
    doc = Document(DAY01)

    # Supprimer ancien bloc II.1.1 / II.1.2 s'il existe déjà (éviter doublons)
    remove_paragraphs_containing(
        doc,
        [
            "II.1.1 Détection",
            "II.1.2 Grounding",
            "Extension du démonstrateur : YOLOv8n",
            "Grounding DINO (IDEA-Research",
            "Comparaison des deux détecteurs",
            "zivid_yolo_openvla.py",
            "test_zivid_groundingDino.py",
            "returnAllPositions.py — capture",
            "test_yolo.py — test isolé",
        ],
    )

    # TOC : ajouter II.1.2 si absent
    for p in doc.paragraphs:
        if p.text.strip() == "   II.1.1  Détection YOLOv8 + coordonnées 3D (jour 07)":
            nxt = p._element.getnext()
            if nxt is None or "II.1.2" not in (Paragraph(nxt, p._parent).text or ""):
                insert_after(
                    p._element,
                    doc,
                    "   II.1.2  Grounding DINO + open-vocabulary (jour 07)",
                )
            break

    # Corps : après paragraphe 224×224
    for p in doc.paragraphs:
        if "Pour OpenVLA, l'image est convertie RGBA → RGB" in p.text:
            anchor = p._element
            blocks = [
                ("Heading 3", "II.1.1 Détection YOLOv8 et coordonnées 3D (jour 07)"),
                (
                    "Normal",
                    "YOLOv8n détecte des objets COCO en 2D ; le centre (u,v) est projeté dans le "
                    "nuage Zivid pour obtenir (X,Y,Z). Le prompt OpenVLA inclut le label et la position : "
                    "pick up the {label} at position X=… Y=… Z=…",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/test/zivid_yolo_openvla.py — pipeline YOLO complet.",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/functions/returnAllPositions.py — capture + YOLO + liste 3D.",
                ),
                ("Heading 3", "II.1.2 Grounding DINO — détection open-vocabulary (jour 07)"),
                (
                    "Normal",
                    "Grounding DINO (grounding-dino-base) accepte une requête texte libre "
                    "(ex. « cell phone. ») et retourne des boîtes sans vocabulaire fixe. Même "
                    "projection 3D que YOLO. Différence clé : le prompt OpenVLA ne contient que "
                    "pick up the {label} — les XYZ restent pour le contrôleur robot, pas pour le VLA.",
                ),
                (
                    "List Bullet",
                    "scripts/openVLA_ZIVID/test/test_zivid_groundingDino.py — pipeline Grounding DINO.",
                ),
                ("Heading 3", "II.1.3 Comparaison YOLO vs Grounding DINO"),
                (
                    "Normal",
                    "YOLO : rapide, classes COCO, XYZ dans le prompt. Grounding DINO : consigne "
                    "naturelle, flexibilité sémantique, XYZ hors prompt OpenVLA. Détail : "
                    "OpenVLA_day07_robot_data.docx.",
                ),
            ]
            for style, text in reversed(blocks):
                anchor = insert_after(anchor, doc, text, style)
            break

    # Figure 2 explication — mention Grounding DINO
    for p in doc.paragraphs:
        if p.text.startswith("Étapes illustrées : (1) capture Zivid"):
            p.text = (
                "Étapes illustrées (pipeline YOLO, Figure 2) : capture Zivid ; détection 2D YOLOv8n ; "
                "pixel (u,v) ; conversion 3D ; OpenVLA (image + prompt avec XYZ). Variante Grounding DINO "
                "(jour 07) : même chaîne 2D→3D, mais détection par texte libre et prompt OpenVLA sans XYZ "
                "(test_zivid_groundingDino.py)."
            )
            break

    # II.5 entrées modèle
    for p in doc.paragraphs:
        if "variante jour 07" in p.text.lower() or p.text.startswith("Entrées modèle : image RGB"):
            p.text = (
                "Entrées modèle : image RGB 224×224 (Zivid) + consigne texte ; jour 07 — variante YOLO "
                "(label + XYZ dans le prompt) ou Grounding DINO (label seul, XYZ pour le robot)."
            )
            break

    # Schéma cible
    for p in doc.paragraphs:
        if p.text.startswith("Schéma cible du démonstrateur"):
            p.text = (
                "Schéma cible : Zivid → détection 2D (YOLO ou Grounding DINO) → coords 3D → "
                "prétraitement image + prompt → OpenVLA → action → UR (RTDE) → nouvelle capture."
            )
            break

    # III.2
    for p in doc.paragraphs:
        if "YOLOv8n relie" in p.text:
            p.text = (
                "Stage : Zivid fournit 2D et 3D ; OpenVLA consomme le RGB 224×224. YOLOv8n (classes COCO) "
                "ou Grounding DINO (texte libre) relient détection 2D au nuage de points (Figure 2). "
                "Voir comparaison jour 07."
            )
            break

    # II.2 pipeline grounding
    for p in doc.paragraphs:
        if "Pipeline YOLO + 3D" in p.text:
            insert_after(
                p._element,
                doc,
                "Pipeline Grounding DINO : scripts/openVLA_ZIVID/test/test_zivid_groundingDino.py.",
                "List Bullet",
            )
            break

    # Pied de page
    for p in doc.paragraphs:
        if "Rapport final" in p.text and "Fabrice Kouonang" in p.text:
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — I à III (jour 07 : YOLO + Grounding DINO)"
            )
            break

    doc.save(DAY01)
    print(f"Updated {DAY01}")


if __name__ == "__main__":
    build_day07()
    update_day01()
