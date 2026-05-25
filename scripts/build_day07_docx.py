#!/usr/bin/env python3
"""Génère / corrige OpenVLA_day07_robot_data.docx et OpenVLA_day01_stage_CCNB.docx."""

from __future__ import annotations

import os
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO = os.path.join(BASE, "media", "ccnb_innov_logo.jpeg")
TEMPLATE = os.path.join(BASE, "OpenVLA_day05_openvla_integration.docx")
DAY07 = os.path.join(BASE, "OpenVLA_day07_robot_data.docx")
DAY01 = os.path.join(BASE, "OpenVLA_day01_stage_CCNB.docx")


def has_image(paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("a:blip")))


def clear_body_keep_logo(doc: Document) -> None:
    """Garde uniquement le premier paragraphe (logo) du template day05."""
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build_day07() -> None:
    """Copie day05 (logo + media embarqués) puis remplace le contenu."""
    shutil.copy2(TEMPLATE, DAY07)
    doc = Document(DAY07)
    clear_body_keep_logo(doc)

    add_para(doc, "Projet OpenVLA — Jour 07 — Données robot et interprétation OpenVLA")
    add_para(doc, "Par : Fabrice Kouonang")
    add_para(doc, "————————————————————————————")
    add_para(doc, "Superviseur direct : Guillaume Batungwanayo")
    add_para(doc, "Date : 25 mai 2026 | Jour : 07")

    add_heading(doc, "Objectif de la journée")
    add_para(
        doc,
        "Comprendre comment les données circulent entre le robot UR, la caméra Zivid et "
        "OpenVLA : ce qui entre réellement dans le modèle, ce que le robot fournit en retour "
        "dans la boucle fermée, et comment le modèle interprète image + consigne pour produire "
        "une action exécutable.",
    )

    add_heading(doc, "Contexte")
    add_para(
        doc,
        "Le jour 06 a validé la boucle complète Zivid → OpenVLA → UR16e (demoTest.py). "
        "Avant d'étendre le démonstrateur, il est essentiel de clarifier une idée souvent "
        "mal comprise : OpenVLA ne « voit » pas directement l'état interne du robot (joints, "
        "couple, RTDE brut). Dans notre pipeline, la perception passe par la caméra ; le robot "
        "intervient après l'inférence pour appliquer le delta d'action prédit.",
    )

    add_heading(doc, "Schéma du flux de données")
    add_bullet(doc, "Entrées OpenVLA : image RGB 224×224 (Zivid) + consigne langage naturel (prompt).")
    add_bullet(doc, "Sortie OpenVLA : vecteur d'action 7D (dx, dy, dz, rx, ry, rz, gripper).")
    add_bullet(
        doc,
        "Données robot (RTDE) : pose TCP actuelle lue après chaque prédiction — non envoyée au modèle.",
    )
    add_bullet(
        doc,
        "Boucle : capture → inférence → lecture pose → moveL + pince → nouvelle capture (vision = feedback).",
    )

    add_heading(doc, "1. Ce qu'OpenVLA reçoit (entrées du modèle)")
    add_heading(doc, "1.1 Image — provenance Zivid, pas du robot", level=2)
    add_para(
        doc,
        "La caméra Zivid capture une frame 2D/3D (RGBA). Le script retire le canal alpha, "
        "convertit en RGB et redimensionne à 224×224 pixels — format attendu par le "
        "préprocesseur Vision-Language du modèle openvla-7b.",
    )
    add_bullet(doc, "Fichier : capture_latest.png (archive locale).")
    add_bullet(doc, "Code : capture_zivid() dans demoTest.py et test_zivid_openvla.py.")

    add_heading(doc, "1.2 Consigne textuelle — prompt VLA", level=2)
    add_para(doc, "OpenVLA attend un prompt structuré, identique à l'entraînement Bridge :")
    add_para(
        doc,
        "In: What action should the robot take to {INSTRUCTION}?\nOut:",
        bold=True,
    )
    add_bullet(doc, "INSTRUCTION exemple : pick up the object on the table.")
    add_bullet(
        doc,
        "AutoProcessor fusionne tokens texte + tenseur image → entrées du réseau sur GPU (bfloat16).",
    )

    add_heading(doc, "1.3 Ce que le modèle ne reçoit pas (dans notre démo)", level=2)
    add_bullet(doc, "Positions articulaires du UR (q_actual).")
    add_bullet(doc, "Pose TCP en entrée du réseau (utilisée seulement après predict_action).")
    add_bullet(doc, "Nuage de points Zivid 3D (PLY) — non consommé par openvla-7b aujourd'hui.")
    add_bullet(doc, "Images embarquées sur le robot (wrist cam) — non utilisées ici.")

    add_heading(doc, "2. Données en provenance du robot UR (RTDE)")
    add_heading(doc, "2.1 Lecture — rtde_receive", level=2)
    add_para(
        doc,
        "RTDEReceiveInterface expose l'état temps réel. Le démonstrateur utilise "
        "getActualTCPPose() : 6 valeurs [x, y, z, rx, ry, rz] en mètres et radians (repère base).",
    )
    add_heading(doc, "2.2 Commande — rtde_control", level=2)
    add_bullet(doc, "moveL(new_pose, speed, acceleration) — déplacement linéaire de l'effecteur.")
    add_bullet(doc, "setToolDigitalOut(0, True/False) — ouverture / fermeture pince UR native.")
    add_heading(doc, "2.3 Rôle du robot dans la boucle", level=2)
    add_para(
        doc,
        "Le robot n'alimente pas l'inférence : il exécute la sortie du modèle. La pose courante "
        "sert de référence pour transformer un delta prédit en pose absolue cible :",
    )
    add_para(
        doc,
        "new_pose[i] = current_pose[i] + action[i] × SCALE   (SCALE = 0,05 → max 5 cm/step)",
        bold=True,
    )

    add_heading(doc, "3. Comment OpenVLA interprète les entrées")
    add_heading(doc, "3.1 Architecture (rappel)", level=2)
    add_bullet(doc, "Encodeur vision (DINOv2) → représentation visuelle.")
    add_bullet(doc, "Projecteur → espace du LLM (Llama-2 7B).")
    add_bullet(doc, "LLM + tête d'action → prédiction du prochain mouvement (pas une trajectoire complète).")

    add_heading(doc, "3.2 predict_action — sortie 7D", level=2)
    add_para(doc, 'Appel : vla.predict_action(**inputs, unnorm_key="bridge_orig", do_sample=False)')
    add_bullet(doc, 'unnorm_key="bridge_orig" : dénormalise selon les statistiques du jeu Bridge.')
    add_bullet(doc, "do_sample=False : action la plus probable (déterministe, reproductible).")
    add_bullet(doc, "dx, dy, dz : déplacement cartésien relatif de l'effecteur.")
    add_bullet(doc, "rx, ry, rz : variation d'orientation.")
    add_bullet(doc, "gripper : < 0,5 → fermer ; ≥ 0,5 → ouvrir.")

    add_heading(doc, "3.3 Interprétation sémantique", level=2)
    add_para(
        doc,
        "Le modèle associe pixels + mots de la consigne à un déplacement appris sur des "
        "démonstrations humaines (imitation learning). Il ne planifie pas explicitement comme "
        "un stack MoveIt : chaque step produit un delta local ; la scène mise à jour par la "
        "caméra guide les steps suivants.",
    )

    add_heading(doc, "4. Tableau récapitulatif")
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    hdr = ["Source", "Donnée", "Entrée OpenVLA ?"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ("Zivid", "RGB 224×224", "Oui"),
        ("Consigne utilisateur", "Texte (prompt)", "Oui"),
        ("UR RTDE", "Pose TCP actuelle", "Non (post-inférence)"),
        ("UR RTDE", "Commandes moveL / DO pince", "Non (exécution)"),
    ]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_heading(doc, "5. Scripts de référence")
    add_bullet(doc, "scripts/integration/test/test_zivid_openvla.py — Zivid + inférence sans robot.")
    add_bullet(doc, "scripts/integration/testUR_ZIVID/demoTest.py — boucle complète avec RTDE.")
    add_bullet(doc, "scripts/integration/test/test_openvla.py — chargement modèle seul.")

    add_heading(doc, "6. Pistes d'évolution")
    add_bullet(doc, "Ajouter la pose TCP ou les joints comme entrée proprioceptive (si le modèle le supporte).")
    add_bullet(doc, "Exploiter le nuage Zivid pour calibration ou contraintes géométriques.")
    add_bullet(doc, "Caméra poignet UR en complément de la vue fixe Zivid.")

    add_heading(doc, "Livrables jour 07")
    add_bullet(doc, "Documentation : OpenVLA_day07_robot_data.docx.")
    add_bullet(doc, "Mise à jour rapport final : section II.5 dans OpenVLA_day01_stage_CCNB.docx.")
    add_bullet(doc, "README — entrée Jour 07.")

    doc.save(DAY07)
    print(f"Created {DAY07} (logo from day05 template)")


def remove_paragraphs_by_text(doc: Document, needles: list[str]) -> None:
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if any(n in text for n in needles):
            p._element.getparent().remove(p._element)


def insert_after_paragraph(doc: Document, anchor_text: str, blocks: list[tuple[str, str]]) -> None:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    anchor = None
    for p in doc.paragraphs:
        if anchor_text in p.text:
            anchor = p._element
            break
    if anchor is None:
        raise RuntimeError(f"Anchor not found: {anchor_text!r}")

    prev = anchor
    for style, text in blocks:
        new_p = OxmlElement("w:p")
        prev.addnext(new_p)
        para = Paragraph(new_p, doc)
        if style and style != "Normal":
            para.style = style
        para.add_run(text)
        prev = new_p


def fix_day01_stage() -> None:
    doc = Document(DAY01)

    # Supprimer l'ancien bloc II.5 mal inséré (ordre inversé, titre en dernier)
    remove_paragraphs_by_text(
        doc,
        [
            "II.5 Flux de données robot",
            "Clarification essentielle pour le démonstrateur",
            "Le robot intervient après predict_action",
            "Entrées modèle : RGB Zivid",
            "Sortie modèle : 7D",
            "Robot : exécution et pose courante",
            "Détail complet : OpenVLA_day07",
        ],
    )

    # Table des matières — ajouter II.5 (ligne indentée, pas le titre Heading du corps)
    for p in doc.paragraphs:
        if p.text.startswith("   II.4  Exécution robotique"):
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph

            if "II.5" in (p._element.getnext().text if p._element.getnext() is not None else ""):
                break
            new_p = OxmlElement("w:p")
            p._element.addnext(new_p)
            para = Paragraph(new_p, doc)
            para.add_run(
                "   II.5  Flux de données robot, Zivid et interprétation OpenVLA (jour 07)"
            )
            break

    # Corps — II.5 après le schéma cible (ordre correct)
    insert_after_paragraph(
        doc,
        "Schéma cible du démonstrateur",
        [
            (
                "Heading 2",
                "II.5 Flux de données robot, Zivid et interprétation OpenVLA (jour 07)",
            ),
            (
                "Normal",
                "Clarification essentielle pour le démonstrateur CCNB-INNOV : OpenVLA reçoit une image "
                "RGB 224×224 issue de la Zivid et une consigne en langage naturel (prompt VLA). Il ne "
                "reçoit pas l'état articulaire ni la pose TCP du UR en entrée du réseau.",
            ),
            (
                "Normal",
                "Le robot intervient après predict_action : RTDE lit getActualTCPPose(), puis le script "
                "calcule new_pose = current_pose + action × SCALE (SCALE = 0,05) et envoie moveL ainsi "
                "que la commande pince (setToolDigitalOut). La boucle visuelle — nouvelle capture Zivid "
                "à chaque itération — constitue le retour capteur principal.",
            ),
            (
                "List Bullet",
                "Entrées modèle : image RGB 224×224 (Zivid) + texte (AutoProcessor → GPU bfloat16).",
            ),
            (
                "List Bullet",
                "Sortie modèle : 7D (dx, dy, dz, rx, ry, rz, gripper), dénormalisation bridge_orig.",
            ),
            (
                "List Bullet",
                "Données robot : pose TCP et commandes RTDE — exécution uniquement, pas d'inférence directe.",
            ),
            (
                "Normal",
                "Rapport détaillé jour 07 : OpenVLA_day07_robot_data.docx. Scripts : demoTest.py, "
                "test_zivid_openvla.py.",
            ),
        ],
    )

    # Pied de page
    for p in doc.paragraphs:
        if "Rapport final" in p.text and "Fabrice Kouonang" in p.text:
            p.text = (
                "Programmeur : Fabrice Kouonang | Projet : OpenVLA — Stage CCNB-INNOV | "
                "Rapport final — Parties I à III (II.5 jour 07 : flux données robot / OpenVLA)"
            )
            break

    doc.save(DAY01)
    print(f"Fixed {DAY01}")


if __name__ == "__main__":
    build_day07()
    fix_day01_stage()
